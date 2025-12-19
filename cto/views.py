# from json.decoder import JSONDecoder
# from typing import Any, Dict
from django.conf.urls import url
from django.urls.conf import path
import docx
import locale
import json
import requests
import urllib
from urllib import request, parse


import io
from operator import index
from pathlib import Path
import os


# Para utilizar algunas de las funciones de la librería
from docx import Document
from docx.shared import Inches, Pt, Cm

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_BREAK_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_TABLE_DIRECTION
from docx.enum.style import WD_STYLE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


from django.shortcuts import render, redirect, get_list_or_404
from django.views import generic
from django.views.generic import TemplateView, ListView, CreateView
from django.contrib.messages.views import SuccessMessageMixin

from django.urls import reverse_lazy
from django.contrib.auth.decorators import login_required, permission_required
from django.contrib.auth.models import User
from django.http import HttpResponse, HttpResponseRedirect, HttpRequest, JsonResponse, HttpResponseServerError

from datetime import datetime, date
from django.contrib import messages
from django.db.models import Q
from django.contrib.auth import authenticate
from requests.api import get
from bases.views import SinPrivilegios
from bases.views import *
from .models import Departamento, Partes, Contratos, Doctos, Tipocontrato, Requisitos, Valida, Secuencia, Regimen, Ciclos, Puestos, Campus
from .forms import DepartamentoForm, PartesForm, ContratosForm, PuestosForm, PartesForm2
from bases.views import SinPrivilegios
from django.core.files.storage import FileSystemStorage, get_storage_class
from api.serializer import TipocontratoSerializer

from django.contrib.auth import update_session_auth_hash
from django.contrib.auth.forms import PasswordChangeForm
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.contrib import messages
from django.http import HttpResponseRedirect
from django.urls import reverse
from babel.dates import format_date


from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from .models import Partes
from datetime import datetime

# --- API auxiliar: campos faltantes del Sujeto de Contrato ---
@login_required(login_url="/login/")
def partes_faltantes(request, id):
    partes = Partes.objects.filter(pk=id).first()
    if not partes:
        return JsonResponse({"error": "Sujeto no encontrado"}, status=404)

    checks = [
        ("rfc", "R.F.C."),
        ("curp", "CURP"),
        ("titulo_profParte", "Título profesional"),
        ("universidadParte", "Universidad"),
        ("cedula_profParte", "Cédula profesional"),
        ("domicilioParte", "Domicilio"),
    ]

    def is_empty(val):
        if val is None:
            return True
        s = str(val).strip()
        # Quitar etiquetas HTML simples del RichText
        try:
            import re
            s = re.sub(r"<[^>]*>", "", s).strip()
        except Exception:
            pass
        return len(s) == 0

    missing = []
    for field, label in checks:
        try:
            val = getattr(partes, field)
        except Exception:
            val = None
        if is_empty(val):
            missing.append({"field": field, "label": label})

    # Beneficiarios obligatorios: al menos uno
    try:
        b1 = getattr(partes, 'beneficiario1', None)
        b2 = getattr(partes, 'beneficiario2', None)
        b3 = getattr(partes, 'beneficiario3', None)
    except Exception:
        b1 = b2 = b3 = None

    if all(is_empty(x) for x in [b1, b2, b3]):
        missing.append({"field": "beneficiarios", "label": "Beneficiarios (al menos uno)"})

    return JsonResponse({"count": len(missing), "missing": missing})

def generar_documento_beneficiarios(request, pk):
    # Obtener la parte registrada
    parte = Partes.objects.get(pk=pk)

    # Crear un documento de Word
    document = Document()

    # Estilo para el título
    title = document.add_heading('ANEXO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True

    # Agregar contenido al documento
    document.add_paragraph('LA DESIGNACIÓN DE BENEFICIARIOS POR PARTE DEL TRABAJADOR(A) PARA DAR CUMPLIMIENTO A LO PREVISTO EN EL ARTÍCULO 25 FRACCIÓN X DE LA LEY FEDERAL DEL TRABAJO Y DE CONFORMIDAD A LO QUE SE ESTABLECE EN EL ARTÍCULO 501 DE LA LEY FEDERAL DEL TRABAJO.')
    document.add_paragraph('DATOS DEL TRABAJADOR')
    document.add_paragraph(f'NOMBRE: {parte.nombresParte} {parte.apellidoPaternoParte} {parte.apellidoMaternoParte}')
    document.add_paragraph(f'CURP: {parte.curp}')
    document.add_paragraph('Por medio del presente escrito manifiesto que es mi deseo designar a mi beneficiario(s), a efecto de que reciban el pago de los salarios y prestaciones devengadas o no cobradas que se generen por mi fallecimiento o desaparición derivada de un acto delincuencial, de acuerdo a lo previsto en el artículo 25 fracción x de la ley federal del trabajo y de conformidad a lo que se establece en el artículo 501 de la ley federal del trabajo.')

    # Tabla de beneficiarios
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Ajustar el ancho de las columnas
    for idx, width in enumerate([Inches(4.0), Inches(2.0), Inches(2.0)]):  # Doble ancho para la primera columna
        col = table.columns[idx]
        col.width = width

    # Encabezados de la tabla
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'NOMBRE'
    hdr_cells[1].text = 'PARENTESCO'
    hdr_cells[2].text = 'PORCENTAJE'

    # Agregar datos de beneficiarios
    beneficiarios = [
        (parte.beneficiario1, parte.parentesco1, parte.porcentaje1),
        (parte.beneficiario2, parte.parentesco2, parte.porcentaje2),
        (parte.beneficiario3, parte.parentesco3, parte.porcentaje3),
    ]

    for beneficiario, parentesco, porcentaje in beneficiarios:
        if beneficiario:  # Solo agregar filas si hay un beneficiario
            row_cells = table.add_row().cells
            row_cells[0].text = beneficiario
            row_cells[1].text = parentesco if parentesco else ""
            row_cells[2].text = porcentaje if porcentaje else ""

    # Obtener la fecha del sistema y el nombre del mes en español
    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
    ]
    fecha_actual = datetime.now()
    dia = fecha_actual.strftime("%d")
    mes = meses[fecha_actual.month - 1]  # Restamos 1 porque los meses en Python van de 0 a 11
    año = fecha_actual.strftime("%Y")

    # Agregar fecha y firma# Agregar tres párrafos vacíos para los saltos de línea
    document.add_paragraph('')
    document.add_paragraph('')


# Agregar el párrafo con el texto

    document.add_paragraph(f'MERIDA, YUCATÁN A {dia} DEL MES DE {mes.upper()} DEL AÑO {año}.').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('__________________________________________').alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('NOMBRE Y FIRMA DEL TRABAJADOR (A)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    

    # Guardar el documento en un buffer
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    # Devolver el documento como respuesta
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=BENEFICIARIOS_{parte.nombresParte}.docx'
    return response
# @login_required
# def change_password(request):
#     if request.method == 'POST':
#         form = PasswordChangeForm(request.user, request.POST)
#         if form.is_valid():
#             user = form.save()
#             update_session_auth_hash(request, user)  # Actualiza la sesión del usuario para mantener la autenticación
#             messages.success(request, 'Tu contraseña ha sido cambiada exitosamente!')
#             return redirect('change_password')
#         else:
#             messages.success(request, 'Ocurrió un error!')
#     else:
#         form = PasswordChangeForm(request.user)
#     return render(request, 'changepass.html', {'form': form})


class VistaBaseCreate(SuccessMessageMixin, SinPrivilegios, generic.CreateView):
    context_object_name = 'obj'
    success_message = "Registro Agregado Satisfactoriamente"

    def form_valid(self, form):
        form.instance.uc = self.request.user
        return super().form_valid(form)


class VistaBaseEdit(SuccessMessageMixin, SinPrivilegios, generic.UpdateView):
    context_object_name = 'obj'
    success_message = "Registro Actualizado Satisfactoriamente"

    def form_valid(self, form):
        #form.instance.um = self.request.user.id
        return super().form_valid(form)


class DepartamentoView(SinPrivilegios, generic.ListView):
    model = Departamento
    template_name = "cto/departamento_list.html"
    context_object_name = "obj"
    permission_required = "cto.view_departamento"


class DepartamentoNew(VistaBaseCreate):
    model = Departamento
    template_name = "cto/departamento_form.html"
    form_class = DepartamentoForm
    success_url = reverse_lazy("cto:departamento_list")
    permission_required = "cto.add_departamento"


class DepartamentoEdit(VistaBaseEdit):
    model = Departamento
    template_name = "cto/departamento_form.html"
    form_class = DepartamentoForm
    success_url = reverse_lazy("cto:departamento_list")
    permission_required = "cto.change_departamento"


@login_required(login_url="/login/")
@permission_required("cto.change_departamento", login_url="/login/")
def departamentoInactivar(request, id):
    departamento = Departamento.objects.filter(pk=id).first()

    if request.method == "POST":
        if departamento:
            departamento.estado = not departamento.estado
            departamento.save()
            return HttpResponse("OK")
        return HttpResponse("FAIL")

    return HttpResponse("FAIL")


class PartesView(SinPrivilegios, generic.ListView):
    model = Partes
    form_class = PartesForm
    template_name = "cto/partes_list.html"
    context_object_name = "obj"
    success_url = reverse_lazy("cto:partes_list")
    permission_required = "cto.view_partes"

    # def get_queryset(self):
    #     current_userx = self.request.user.id
    #     # print(current_userx)
    #     queryset = Partes.objects.filter(user_id=current_userx)
    #     xdepa = 100
    #     for part in queryset:
    #         xdepa = part.claveDepartamento_id

    #     # print(xdepa)

    #     querydep = Departamento.objects.all()
    #     # print(querydep)
    #     for depa in querydep:

    #         if depa.claveDepartamento == xdepa:
    #             # print(depa.claveDepartamento)
    #             xr1 = "0"+str(depa.rango1)
    #             xr2 = "0"+str(depa.rango2)
    #             # print(xdepa)
    #             # print(xr1)
    #             # print(xr2)
    #             return Partes.objects.filter(Q(claveDepartamento__gte=xr1),  Q(claveDepartamento__lte=xr2))
    #         # else:
    #         #    messages.ERROR (request, '¡Rango de Departamentos no registrado!')
    #         #    return HttpResponseRedirect('/admn/')

    def get_context_data(self, **kwargs):
        # Call the base implementation first to get a context
        context = super(PartesView, self).get_context_data(**kwargs)
        # Get the blog from id and add it to the context
        return context


class PartesNew(VistaBaseCreate):
    model = Partes
    template_name = "cto/partes_form.html"
    form_class = PartesForm
    success_url = reverse_lazy("cto:partes_list")
    permission_required = "cto.add_partes"


class PartesEdit(VistaBaseEdit):
    model = Partes
    template_name = "cto/partes_form.html"
    form_class = PartesForm2
    success_url = reverse_lazy("cto:partes_list")
    permission_required = "cto.change_partes"


@login_required(login_url="/login/")
@permission_required("cto.change_partes", login_url="/login/")
def partesInactivar(request, id):
    partes = Partes.objects.filter(pk=id).first()

    if request.method == "POST":
        if partes:
            partes.estado = not partes.estado
            partes.save()
            return HttpResponse("OK")
        return HttpResponse("FAIL")

    return HttpResponse("FAIL")


class ContratosView(SinPrivilegios, generic.ListView):
    model = Contratos
    template_name = "cto/contrato_list.html"
    context_object_name = "obj"
    success_url = reverse_lazy("cto:contrato_list")
    permission_required = "cto.view_contratos"

    def get_queryset(self):
        xtipo = 0
        tipocontrato = Tipocontrato.objects.filter(marcatipoContrato=True)
        for x in tipocontrato:
            if x.marcatipoContrato == True:
                xtipo = x.id
                # print(xtipo)
        # r = requests.get('http://127.0.0.1:8000/api/v1/tipocontrato/7')
        # print (r.content)
        # print (r.status_code)
        # print (r.headers)
        # print (r.json)
        # x=urllib.request.urlretrieve('http://127.0.0.1:8000/api/v1/tipocontrato/7')
        # print(x)8

        current_userx = self.request.user.id
        #conditions = dict(current_user=current_userx, uc_id=self.request.user)
        #queryset = queryset.filter(**conditions)
        # return Contratos.objects.all()

        return Contratos.objects.filter(
            (Q(current_user=current_userx) | Q(uc_id=self.request.user)), Q(
                tipocontrato_id=xtipo)
        )
        # return SpyorEnc.objects.filter(
        #    Q(current_user=current_userx) | Q(uc_id=self.request.user) | Q(el_jefe=current_userx)
        # )

    def get_context_data(self, **kwargs):

        # Call the base implementation first to get a context
        context = super(ContratosView, self).get_context_data(**kwargs)
        # Get the blog from id and add it to the context
        context['some_data'] = Partes.objects.all()
        context['some_data2'] = Departamento.objects.all()
        context['some_data3'] = Tipocontrato.objects.filter(marcatipoContrato=True).order_by('-fm')  # Ordena los datos por la fecha de actualización
        return context


class ContratosView2(SinPrivilegios, generic.ListView):
    model = Contratos
    template_name = "cto/contrato_list.html"
    context_object_name = "obj"
    success_url = reverse_lazy("cto:contrato_list")
    permission_required = "cto.view_contratos"

    def get_queryset(self):
        xtipo = 0
        tipocontrato = Tipocontrato.objects.filter(marcatipoContrato=True)
        for x in tipocontrato:
            if x.marcatipoContrato == True:
                xtipo = x.id
                # print(xtipo)
        # r = requests.get('http://127.0.0.1:8000/api/v1/tipocontrato/7')
        # print (r.content)
        # print (r.status_code)
        # print (r.headers)
        # print (r.json)
        # x=urllib.request.urlretrieve('http://127.0.0.1:8000/api/v1/tipocontrato/7')
        # print(x)8

        current_userx = self.request.user.id
        #conditions = dict(current_user=current_userx, uc_id=self.request.user)
        #queryset = queryset.filter(**conditions)
        # return Contratos.objects.all()

        return Contratos.objects.filter(
            (Q(current_user=current_userx) | Q(uc_id=self.request.user)), Q(
                tipocontrato_id=xtipo)
        )
        # return SpyorEnc.objects.filter(
        #    Q(current_user=current_userx) | Q(uc_id=self.request.user) | Q(el_jefe=current_userx)
        # )

    def get_context_data(self, **kwargs):

        # Call the base implementation first to get a context
        context = super(ContratosView, self).get_context_data(**kwargs)
        # Get the blog from id and add it to the context
        context['some_data'] = Partes.objects.all()
        context['some_data2'] = Departamento.objects.all()
        context['some_data3'] = Tipocontrato.objects.filter(marcatipoContrato=True).order_by('-fm')  # Ordena los datos por la fecha de actualización
        return context


@login_required(login_url='/login/')
@permission_required('cto.add_contratos', login_url='bases:sin_privilegios')
def contratos2(request, contrato_id=None):

    template_name = 'cto/contrato.html'
    detalle = {}
    #secuencia_data = {}
    #xUsuario = (request.user.id)
    # print(xUsuario)
    #print (contrato_id)

    if contrato_id:
        contrato = Contratos.objects.filter(estado=True, id=contrato_id)
        c = contrato.first()
        xUsuario = (c.parte2_id)
        # print(xUsuario)
        partes = Partes.objects.filter(estado=True, id=xUsuario)
        p = partes.first()
        #print (p)
        dx = p.claveDepartamento_id
    else:
        xUsuario = (request.user.username)
        xUsuario2 = (request.user.id)
        # print(xUsuario2)
        partes = Partes.objects.filter(estado=True, usuario=xUsuario)
        p = partes.first()
        # Fallback: algunos usuarios no tienen vinculado el OneToOne o el campo 'usuario'
        if not p:
            partes = Partes.objects.filter(estado=True, user_id=xUsuario2)
            p = partes.first()
        if not p:
            messages.error(request, 'No se encontró un "Sujeto de contrato" asociado a tu usuario. Asigna tu usuario en Catálogo > Partes.')
            return HttpResponseRedirect('/cto/contratos/')
        #print (p)
        dx = p.claveDepartamento_id

    #d1 = p.claveDepartamento

    #rfcparte = (p.rfc)
    # print(rfcparte)

    # print(dx)

    requisitos = Requisitos.objects.filter(estado=True)
    departamentos = Departamento.objects.filter(
        estado=True, claveDepartamento=dx)
    departamentos2 = Departamento.objects.filter(estado=True)
    

    d2 = departamentos.first()
    if not d2:
        messages.error(request, f'Departamento no encontrado para la clave "{dx}". Verifica el catálogo de Departamentos y la clave en tu registro de Sujeto.')
        return HttpResponseRedirect('/cto/contratos/')
   
    d3 = (d2.claveDepartamento)
    #print(d3)
    #r1 = (d2.rango1)
    #r2 = (d2.rango2)

    # Asegurar padding (comparación lexicográfica consistente contra claves de 3 dígitos con ceros a la izquierda)
    try:
        r1 = f"{int(d2.rango1):03d}" if d2.rango1 is not None else "000"
    except Exception:
        r1 = "000"
    try:
        r2 = f"{int(d2.rango2):03d}" if d2.rango2 is not None else "999"
    except Exception:
        r2 = "999"

    #print(d2.rango1)
    #print(d2.rango2)

    if d2.rango1 == 2:
        r1 = "002"
        r2 = "048"
        
    if d2.rango1 == 14:
        r1 = "014"
        r2 = "014"  
  
    if d2.rango1 == 40:
        r1 = "040"
        r2 = "040"
        
    if d2.rango1 == 29:
        r1 = "029"
        r2 = "029"
        
    if d2.rango1 == 41:
        r1 = "041"
        r2 = "041"     
        
    if d2.rango1 == 42:
        r1 = "042"
        r2 = "042"
        
    if d2.rango1 == 43:
        r1 = "043"
        r2 = "043"
        
    if d2.rango1 == 44:
        r1 = "044"
        r2 = "044"            
        

    secuencia = (d2.f001)

    # print(d2.testigoUsual1)
    # print(d2.testigoUsual2)
    #print (d3)
    # print(secuencia)
    # if r1 == "002":
    #     r1 = "0"

    # print (r1)
    # print (r2)

    partes3 = Partes.objects.filter(Q(estado=True),  Q(claveDepartamento_id__gte=r1), Q(
        claveDepartamento_id__lte=r2)).order_by('nombreParte')

    # print (partes3)

    partes2 = Partes.objects.filter(estado=True).order_by('nombreParte')
    # print (partes2)

    if secuencia:

        r1 = int(secuencia[:5])
        f1 = secuencia[5:8]

        r2 = int(secuencia[8:13])
        f2 = secuencia[13:16]

        r3 = int(secuencia[16:21])
        f3 = secuencia[21:24]

        # Elegir el responsable (DIC) desde la secuencia
        dic_user_id = None
        if f1 == 'DIC':
            dic_user_id = r1
        elif f2 == 'DIC':
            dic_user_id = r2
        elif f3 == 'DIC':
            dic_user_id = r3

        if dic_user_id is None:
            messages.error(request, 'Secuencia administrativa inválida: no se encontró responsable DIC (f1/f2/f3).')
            return HttpResponseRedirect('/cto/contratos/')

        funcionario = Partes.objects.filter(user_id=dic_user_id)
        a2 = funcionario.first()
        if not a2:
            messages.error(request, 'Secuencia administrativa inválida: el usuario DIC no está dado de alta en Partes.')
            return HttpResponseRedirect('/cto/contratos/')
        a3 = a2.id
        fun = Partes.objects.get(pk=a3)

        r4 = int(secuencia[24:29])
        f4 = secuencia[29:32]

        r5 = int(secuencia[32:37])
        f5 = secuencia[37:40]

        r6 = int(secuencia[40:45])
        f6 = secuencia[45:48]

        # print(r1)
        # print(f1)
        # print(r2)
        # print(f2)
        # print(r3)
        # print(f3)
        # print(r4)
        # print(f4)
        # print(r5)
        # print(f5)
        # print(r6)
        # print(f6)
    else:
        #print ("secuencia no asignada")
        messages.error(request, '¡Secuencia administrativa no asignada!')
        #messages.error = ('secuencia no asignada')
        return HttpResponseRedirect('/cto/contratos/')

    if request.method == "GET":
        enc = Contratos.objects.filter(pk=contrato_id).first()
        tipocontratox = Tipocontrato.objects.filter(marcatipoContrato=True)
        tipocontratox = tipocontratox.first()

        tipocontrato = tipocontratox.id
        if not enc:
            encabezado = {
                'id': "",
                'uc_id': "",
                'tipocontrato': tipocontrato,
                'datecontrato': datetime.today(),
                'datecontrato_ini': "",
                'datecontrato_fin': "",
                'parte1': 164,
                'enCalidadDe1': '"CLIENTE"',
                'parte2': "",
                'enCalidadDe2': '"PRESTADOR DE SERVICIOS"',
                'lugarContrato': "",
                'ciudadContrato': "Mérida",
                'estadoContrato': "Yucatán",
                'paisContrato': "México",
                'importeContrato': 0.00,
                'npContrato': 1,

                'empresa': "",
                'tipo_sociedad': "",
                'objeto_social': "",
                'registro_mercantil': "",
                'fecha_constitucion': "",
                'rfc_sociedad': "",
                'domicilio_sociedad': "",
                'testimonio': "",
                'clausula': "",
                
                
                'marcaVehiculo': "",
                'modeloVehiculo': "",
                'tipoVehiculo': "",
                'motorVehiculo': "",
                'serieVehiculo': "",
                'placasVehiculo': "",
                'facturaVehiculo': "",
                'fechaFactura': "",
                'expedidaFactura': "",
                'tarjetaVehiculo': "",
                'fechaTarjeta': "",
                'expedidaTarjeta': "",
                'polizaVehiculo': "",
                'fechaPoliza': "",
                'expedidaPoliza': "",


                'imppContrato': 0.00,
                'vhppContrato': 285.00,
                'totalhorasContrato': "",
                
                'testigoContrato1': d2.testigoUsual1,
                'testigoContrato2': d2.testigoUsual2,
                
                # 'testigoContrato1': "",
                # 'testigoContrato2': "",
                
                'versionContrato': "",

                'status': "CAP",
                'rcap': xUsuario2,


                'rstep1': r1,
                'rstep2': r2,
                'rstep3': r3,
                'rstep4': r4,
                'rstep5': r5,
                'rstep6': r6,

                'astep1': f1,
                'astep2': f2,
                'astep3': f3,
                'astep4': f4,
                'astep5': f5,
                'astep6': f6,
                'devuelto_por': "",

            }
            detalle = None
        else:

            encabezado = {
                'id': enc.id,
                'uc_id': enc.uc_id,
                'tipocontrato': enc.tipocontrato,
                'datecontrato': enc.datecontrato,
                'datecontrato_ini': enc.datecontrato_ini,
                'datecontrato_fin': enc.datecontrato_fin,
                'parte1': enc.parte1,
                'enCalidadDe1': enc.enCalidadDe1,
                'parte2': enc.parte2,
                'enCalidadDe2': enc.enCalidadDe2,
                'lugarContrato': enc.lugarContrato,
                'ciudadContrato': enc.ciudadContrato,
                'estadoContrato': enc.estadoContrato,
                'paisContrato': enc.paisContrato,
                'importeContrato': enc.importeContrato,
                'npContrato': enc.npContrato,
                
                
                'empresa': enc.empresa,
                'tipo_sociedad': enc.tipo_sociedad,
                'objeto_social': enc.objeto_social,
                'registro_mercantil': enc.registro_mercantil,
                'fecha_constitucion': enc.fecha_constitucion,
                'rfc_sociedad': enc.rfc_sociedad,
                'domicilio_sociedad': enc.domicilio_sociedad,
                'testimonio': enc.testimonio,
                'clausula': enc.clausula,
                


                'marcaVehiculo': enc.marcaVehiculo,
                'modeloVehiculo': enc.modeloVehiculo,
                'tipoVehiculo': enc.tipoVehiculo,
                'motorVehiculo': enc.motorVehiculo,
                'serieVehiculo': enc.serieVehiculo,
                'placasVehiculo': enc.placasVehiculo,
                'facturaVehiculo': enc.facturaVehiculo,
                'fechaFactura': enc.fechaFactura,
                'expedidaFactura': enc.expedidaFactura,
                'tarjetaVehiculo': enc.tarjetaVehiculo,
                'fechaTarjeta': enc.fechaTarjeta,
                'expedidaTarjeta': enc.expedidaTarjeta,
                'polizaVehiculo': enc.polizaVehiculo,
                'fechaPoliza': enc.fechaPoliza,
                'expedidaPoliza': enc.expedidaPoliza,


                'imppContrato': enc.imppContrato,
                'vhppContrato': enc.vhppContrato,
                'totalhorasContrato': enc.totalhorasContrato,
                'testigoContrato1': enc.testigoContrato1,
                'testigoContrato2': enc.testigoContrato2,
                'versionContrato': enc.versionContrato,

                'status': enc.status,
                'rcap': enc.rcap,


                'rstep1': enc.rstep1,
                'rstep2': enc.rstep2,
                'rstep3': enc.rstep3,
                'rstep4': enc.rstep4,
                'rstep5': enc.rstep5,
                'rstep6': enc.rstep6,

                'astep1': enc.astep1,
                'astep2': enc.astep2,
                'astep3': enc.astep3,
                'astep4': enc.astep4,
                'astep5': enc.astep5,
                'astep6': enc.astep6,

                'fstep1': enc.fstep1,
                'fstep2': enc.fstep2,
                'fstep3': enc.fstep3,
                'fstep4': enc.fstep4,
                'fstep5': enc.fstep5,
                'fstep6': enc.fstep6,

                'cstep1': enc.cstep1,
                'cstep2': enc.cstep2,
                'cstep3': enc.cstep3,
                'cstep4': enc.cstep4,
                'cstep5': enc.cstep5,
                'cstep6': enc.cstep6,

                'devuelto_por': enc.devuelto_por,
                'current_user': enc.current_user,

            }

        detalle = Doctos.objects.filter(contrato=enc)

        contexto = {"requi": requisitos, "fun": funcionario, "fun2": partes2, "fun3": partes3, "enc": encabezado, "det": detalle,
                    "departamentos": departamentos, "funcionarios": partes, "departamentos2": departamentos2, "tipocont": tipocontratox, }

        return render(request, template_name, contexto)

    if request.method == "POST":

        tipocontratox = Tipocontrato.objects.filter(marcatipoContrato=True)
        tipocontratox = tipocontratox.first()

        tipocontrato = tipocontratox.id
        # print(tipocontrato)
        datecontrato = request.POST.get("datecontrato")
        datecontrato_ini = request.POST.get("enc_datecontrato_ini")
        datecontrato_fin = request.POST.get("enc_datecontrato_fin")

        parte1 = 164
        parte2 = request.POST.get("enc_nombreParte")

        enCalidadDe1 = tipocontratox.enCalidadDe1
        enCalidadDe2 = tipocontratox.enCalidadDe2

        lugarContrato = request.POST.get("lugarContrato")
        ciudadContrato = "Mérida"
        estadoContrato = "Yucatán"
        paisContrato = "México"
        importeContrato = request.POST.get("enc_importeContrato")
        npContrato = request.POST.get("enc_npContrato")

        empresa = request.POST.get("enc_empresa")
        tipo_sociedad = request.POST.get("enc_tipo_sociedad")
        objeto_social = request.POST.get("enc_objeto_social")
        registro_mercantil = request.POST.get("enc_registro_mercantil")
        fecha_constitucion = request.POST.get("enc_fecha_constitucion")
        rfc_sociedad = request.POST.get("enc_rfc_sociedad")
        domicilio_sociedad = request.POST.get("enc_domicilio_sociedad")
        testimonio = request.POST.get("enc_testimonio")
        clausula = request.POST.get("enc_clausula")
        
        
        
        
        marcaVehiculo = request.POST.get("enc_marcaVehiculo")
        modeloVehiculo = request.POST.get("enc_modeloVehiculo")
        tipoVehiculo = request.POST.get("enc_tipoVehiculo")
        motorVehiculo = request.POST.get("enc_motorVehiculo")
        serieVehiculo = request.POST.get("enc_serieVehiculo")
        placasVehiculo = request.POST.get("enc_placasVehiculo")
        facturaVehiculo = request.POST.get("enc_facturaVehiculo")
        fechaFactura = request.POST.get("enc_fechaFactura")
        expedidaFactura = request.POST.get("enc_expedidaFactura")
        tarjetaVehiculo = request.POST.get("enc_tarjetaVehiculo")
        fechaTarjeta = request.POST.get("enc_fechaTarjeta")
        expedidaTarjeta = request.POST.get("enc_expedidaTarjeta")
        polizaVehiculo = request.POST.get("enc_polizaVehiculo")
        fechaPoliza = request.POST.get("enc_fechaPoliza")
        expedidaPoliza = request.POST.get("enc_expedidaPoliza")

        imppContrato = request.POST.get("enc_imppContrato")

        vhppContrato = 285
        totalhorasContrato = request.POST.get("enc_totalhorasContrato")
        
        testigoContrato1 = request.POST.get("enc_testigoContrato1")
        testigoContrato2 = request.POST.get("enc_testigoContrato2")
        # Fallback a testigos usuales del departamento si no vienen en el POST
        if not testigoContrato1:
            testigoContrato1 = d2.testigoUsual1
        if not testigoContrato2:
            testigoContrato2 = d2.testigoUsual2
            
        # testigoContrato1 = d2.testigoUsual1
        # testigoContrato2 = d2.testigoUsual2
        versionContrato = request.POST.get("versionContrato")
        status = request.POST.get("status")

        fun = Partes.objects.get(pk=a3)
        tip = Tipocontrato.objects.get(pk=tipocontrato)

        if parte2:
            suj = Partes.objects.get(pk=parte2)

        if not contrato_id:
            enc = Contratos(
                tipocontrato=tip,
                datecontrato=datecontrato,
                datecontrato_ini=datecontrato_ini,
                datecontrato_fin=datecontrato_fin,

                parte1=parte1,
                parte2=suj,
                ciudadContrato=ciudadContrato,
                estadoContrato=estadoContrato,
                paisContrato=paisContrato,

                enCalidadDe1=enCalidadDe1,
                enCalidadDe2=enCalidadDe2,

                importeContrato=importeContrato,
                npContrato=npContrato,

                empresa=empresa,
                tipo_sociedad=tipo_sociedad,
                objeto_social=objeto_social,
                registro_mercantil = registro_mercantil,
                fecha_constitucion = fecha_constitucion,
                rfc_sociedad = rfc_sociedad,
                domicilio_sociedad = domicilio_sociedad,
                testimonio = testimonio,
                clausula = clausula,
                


                marcaVehiculo=marcaVehiculo,
                modeloVehiculo=modeloVehiculo,
                tipoVehiculo=tipoVehiculo,
                motorVehiculo=motorVehiculo,
                serieVehiculo=serieVehiculo,
                placasVehiculo=placasVehiculo,
                facturaVehiculo=facturaVehiculo,
                tarjetaVehiculo=tarjetaVehiculo,
                polizaVehiculo=polizaVehiculo,

                fechaFactura=fechaFactura,
                expedidaFactura=expedidaFactura,

                fechaTarjeta=fechaTarjeta,
                expedidaTarjeta=expedidaTarjeta,

                fechaPoliza=fechaPoliza,
                expedidaPoliza=expedidaPoliza,





                imppContrato=imppContrato,
                vhppContrato=vhppContrato,
                totalhorasContrato=totalhorasContrato,
                # Testigos editables por contrato (prefill con valores usuales)
                testigoContrato1=testigoContrato1,
                testigoContrato2=testigoContrato2,

                # testigoContrato1=testigoContrato1,
                # testigoContrato2=testigoContrato2,

                status=status,
                rcap=xUsuario2,
                current_user=xUsuario2,
                rstep1=r1,
                rstep2=r2,
                rstep3=r3,
                rstep4=r4,
                rstep5=r5,
                rstep6=r6,

                astep1=f1,
                astep2=f2,
                astep3=f3,
                astep4=f4,
                astep5=f5,
                astep6=f6,
                devuelto_por=False,
            )
            if enc:
                enc.save()
                data = {
                  "Captura": xUsuario,
                  "id":enc.id,
                  "Sujeto":suj.nombreParte,
                  "Fecha inicial": str(datecontrato_ini),
                  "Fecha final": str(datecontrato_fin) if datecontrato_fin is not None else 'N/A',
                  "Departamento": str(suj.claveDepartamento),
                  "Título": suj.titulo_profParte,
                  "Cédula Prof.": suj.cedula_profParte,   
                  "Tipo de contrato": tip.tipoContrato,
                  "Importe": "${:,.2f}".format(float(importeContrato)) if importeContrato is not None else 'N/A',
                  "Horas": totalhorasContrato,
                  "Actividades": suj.actividadesParte,             
                }
    
                send_telegram_message(request, data)
                # URL del Webhook proporcionada por Make
                webhook_url = 'https://hook.us1.make.com/7vc5sst9w5lfl2zowlf1p4ndf7o6mu7c'
    
                # Enviar solicitud POST al Webhook
                response = requests.post(webhook_url, json=data)
                
                
                
                
                contrato_id = enc.id
                messages.success( request, '¡Contrato Registrado!  Descarga del Contrato disponible')
                return redirect("cto:contrato_edit", contrato_id=contrato_id)
                
        else:
            enc = Contratos.objects.filter(pk=contrato_id).first()
            if enc:

                enc.funcionario = fun
                # Actualiza testigos si vienen en el POST
                if "enc_testigoContrato1" in request.POST:
                    enc.testigoContrato1 = request.POST.get("enc_testigoContrato1")
                if "enc_testigoContrato2" in request.POST:
                    enc.testigoContrato2 = request.POST.get("enc_testigoContrato2")

                enc.save()

        if not id:
            messages.error(
                request, 'No Puedo Continuar No Pude Detectar No. de Contrato')
            return redirect("cto:contrato_list")

    documento = request.POST.get("documento")
    comentarioDocto = request.POST.get("comentarioDocto")
    type(documento)
    req = Requisitos.objects.get(pk=documento)
    vigenciaFinDocto = request.POST.get("enc_vigenciaFinDocto")

    pdf = request.POST.get('pdf2')

    uploaded_file = request.FILES['pdf']
        # print(uploaded_file)

    fs = FileSystemStorage()
        # print(fs)
    name = fs.save(uploaded_file.name, uploaded_file)
        # print(name)
    pdf = name

    if vigenciaFinDocto != "":

        det = Doctos(
            contrato=enc,
            documento=req,
            comentarioDocto=comentarioDocto,
            pdf=pdf,
            vigenciaFinDocto=vigenciaFinDocto,

        )

    if vigenciaFinDocto == "":

        det = Doctos(
            contrato=enc,
            documento=req,
            comentarioDocto=comentarioDocto,
            pdf=pdf,


        )

    if det:
        det.save()

        return redirect("cto:contrato_edit", contrato_id=contrato_id)

    return render(request, template_name, contexto)


class ContratosEdit(VistaBaseEdit):
    model = Contratos
    template_name = "cto/contrato_form.html"
    form_class = ContratosForm
    success_url = reverse_lazy('cto:contrato_list')
    permission_required = "cto.change_contratos"
    context_object_name = 'obj'


class ContratosEdit3(VistaBaseEdit):
    model = Contratos
    template_name = "cto/contrato_form3.html"
    form_class = ContratosForm
    permission_required = "cto.change_contratos"
    context_object_name = 'obj'

    def get_success_url(self):
        return reverse_lazy('contrato_edit', kwargs={'contrato_id': self.object.pk})


@login_required(login_url="/login/")
@permission_required("cto.change_contratos", login_url="/login/")
def coverletter_export(request, id):

    contratos = Contratos.objects.filter(pk=id).first()  # Contrato en curso
    # Información del tipo de contrato
    tipoc = Tipocontrato.objects.get(id=contratos.tipocontrato_id)
    # print(tipoc.id)
    # Validacion de información completa
    valic = Valida.objects.filter(tipocontrato_id=tipoc.id)
    # Datos del contratado **sujeto del contrato
    partes = Partes.objects.get(id=contratos.parte2_id)
    patron = Partes.objects.get(id=164)  # Datos del contratante
    secue = Secuencia.objects.filter(
        tipocontrato_id=tipoc.id).first()  # Primer parrafo del contrato
    ciclo = Ciclos.objects.filter(
        ciclo_actual=True).first()  # Ciclo escolar actual
    # print(tipoc.id)
    # print(secue)
    # Régimen fiscal del contratado
    departamento = Departamento.objects.filter(claveDepartamento=partes.claveDepartamento_id).first()
    # Determinar campus por rango de clave de departamento (robusto ante valores no numéricos)
    try:
        dept_num = int(departamento.claveDepartamento) if departamento and departamento.claveDepartamento else None
    except Exception:
        dept_num = None

    if dept_num is not None and dept_num < 100:
        campus=Campus.objects.filter(pk=2).first()

    if dept_num is not None and dept_num < 300 and dept_num > 99:
        campus=Campus.objects.filter(pk=3).first()

    if dept_num is not None and dept_num < 500 and dept_num > 299:
        campus=Campus.objects.filter(pk=4).first()

    if dept_num is not None and dept_num < 700 and dept_num > 499:
        campus=Campus.objects.filter(pk=5).first()

    # print(campus.direccionCampus)

    regimen = Regimen.objects.get(id=partes.regfiscalParte_id)
    replegal = Partes.objects.get(id=165)  # Datos del contratante

    datecontratox = contratos.datecontrato_ini 
    if contratos.datecontrato_ini.date() > date.today():
        datecontratox = date.today()

    if contratos.importeContrato:
        letras = numero_to_letras(contratos.importeContrato.amount)
    else:
        letras = ""

    if contratos.imppContrato:
        pagolet = numero_to_letras(contratos.imppContrato.amount)
        currency2 = "${:,.2f}".format(contratos.imppContrato.amount)
    else:
        pagolet = ""
        currency2 = ""
    if contratos.importeContrato:
        currency = "${:,.2f}".format(contratos.importeContrato.amount)
    else:
        currency = ""
    document = Document()

    # locale.setlocale(locale.LC_ALL, "es-MX")

    puesto = Puestos.objects.filter(nombrePuesto=partes.clavePuesto).first()

    # footer_para.paragraph_format.page_break_before = True

    # set up font
    font = document.styles['Normal'].font
    font.name = 'Calibri'
    font.bold = True
    font.size = Pt(16)

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style(
        'CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Calibri'

    # set up margins
    sections = document.sections
    for section in sections:

        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.header_distance = Cm(0)

    header = document.sections[0].header
    htable = header.add_table(1, 2, Inches(5))
    # htable.style = "TableGrid"

    for row in htable.rows:
        row.height = Inches(1.0)

    htable.alignment = WD_TABLE_ALIGNMENT.LEFT
    htab_cells = htable.rows[0].cells

    ht0 = htab_cells[0].paragraphs[0]
    paragraph_format = ht0.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.left_indent = Pt(0)

    kh = ht0.add_run()

    THIS_FOLDER = os.path.dirname(os.path.abspath(__file__))
    # print(THIS_FOLDER)
    my_file = os.path.join(THIS_FOLDER, 'logo.png')

    kh.add_picture(my_file, width=Inches(1.00))
    ht0.alignment = WD_ALIGN_PARAGRAPH.LEFT

    ht1 = htab_cells[1].add_paragraph("ESCUELA MODELO, S.C.P.")
    htable.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    htable.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Get the user's fullname
    # if request.user.get_full_name():
    #   document_data_full_name = request.user.get_full_name()
    # else:
    #    document_data_full_name = "[NOMBRE] [APELLIDOS]"

    p001 = document.add_paragraph()
    p001.add_run(tipoc.tituloContrato, style='CommentsStyle')
    p001.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p002 = document.add_paragraph()

    textox = tipoc.textoinicialContrato

    if partes.tituloParte:
        xnombreParte = partes.tituloParte + " " + partes.nombreParte
    else:
        if partes.personaParte == 1:
            xnombreParte = "**********" + " " + partes.nombreParte
        else:
            xnombreParte = partes.nombreParte
        # messages.info(request, message="Registrar: Título del Sujeto del Contrato")

    if partes.curp:
        xcurp = partes.curp[10:11]
    else:
        xcurp = "X"

    if xcurp == "H":
        xelolaParte = "EL"
        xenCalidadDe2 = tipoc.enCalidadDe2
        yenCalidadDe2 = tipoc.enCalidadDe2

    else:
        xelolaParte = "LA"
        xenCalidadDe2 = tipoc.enCalidadDe2f
        yenCalidadDe2 = tipoc.enCalidadDe2f

    xenCalidadDe2 = textox.replace("@enCalidadDe2", xenCalidadDe2)

    textox = xenCalidadDe2
    xenCalidadDe1 = textox.replace("@enCalidadDe1", tipoc.enCalidadDe1)
    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace("@elolaParte", xelolaParte)
    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace("@nombreParte2", xnombreParte)
    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace("@nombreParte", patron.nombreParte)
    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace("@tituloParteRL", replegal.tituloParte)
    textox = xenCalidadDe1
    if contratos.empresa:
        xenCalidadDe1 = textox.replace("@empresa", contratos.empresa)

    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace(
        "@idrep_legalParte", replegal.tituloParte + " " + replegal.nombreParte)

    p002.add_run(xenCalidadDe1, style='CommentsStyle').bold = False
    p002.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if tipoc.id == 1 or tipoc.id == 2 or tipoc.id == 3:
        p003 = document.add_paragraph()
        p003.add_run("CLÁUSULAS", style='CommentsStyle').bold = True
        p003.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if tipoc.id == 7 or tipoc.id == 8 or tipoc.id == 4 or tipoc.id == 14 or tipoc.id == 5 or tipoc.id == 11 or tipoc.id == 13:
        p003 = document.add_paragraph()
        p003.add_run("DECLARACIONES", style='CommentsStyle').bold = True
        p003.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # print(secue.id)
    # secue = Secuencia.objects.get(id=1)
    p004 = document.add_paragraph()
    textosecuex = secue.identificador + ".- " + secue.textoSecuencia
    textosecuex = textosecuex.replace("@enCalidadDe2", yenCalidadDe2)
    # print(yenCalidadDe2)
    textosecuex = textosecuex.replace("@enCalidadDe1", tipoc.enCalidadDe1)
    # print(textosecuex)
    p004.add_run(textosecuex, style='CommentsStyle').bold = True
    paragraph_format = p004.paragraph_format
    paragraph_format.left_indent = Inches(0.0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if tipoc.id == 1:
        nums = (2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19,
                20, 21, 22, 23, 24, 25, 26, 27, 28, 29, 30, 31, 32, 33, 34, 35, 36)

    if tipoc.id == 7:
        nums = (314, 315, 316, 317, 859, 318, 319, 320, 321, 322, 323, 860, 861, 862, 863, 324, 325, 326, 327, 826, 827, 828, 829, 830, 831, 832, 833, 834, 835, 836, 328, 338, 339, 340, 341, 342, 343, 864, 865, 344, 345, 346,
                347, 348, 349, 350, 352, 353, 354, 355, 356, 357, 358, 359, 360, 361, 362, 363, 364, 365, 366, 367, 368, 369, 370, 371, 372, 373, 374, 375, 376, 866, 867, 377, 378, 379, 380, 381)

    if tipoc.id == 8:
        nums = (384, 385, 386, 387, 388, 389, 390, 391, 392, 393, 394, 395, 396, 397, 398, 399, 401, 402, 403, 404, 405, 406, 407, 824, 825, 409, 410, 411, 412, 417, 418, 419, 421,
                422, 423, 424, 425, 426, 427, 428,   429, 430, 431, 432, 433,  434, 435, 436, 437, 438, 439, 440, 441, 442, 443, 444, 445, 446, 447, 448, 842, 843, 845, 846, 449, 450,  453, 454, 455, 456, 457)

    if tipoc.id == 9:
        nums = (460, 461, 462, 463, 464, 465, 466, 467, 468, 469, 470, 471, 472, 473, 474, 475, 476, 478, 479, 480, 481, 482, 483, 484, 822, 823, 485, 486, 487, 488, 489, 490, 491, 492, 493, 494, 495,
                496, 497, 498, 499, 500, 501, 502, 503, 504, 505, 506, 507, 508, 509, 510, 511, 512, 513, 514, 515, 516, 517, 518, 519, 520, 521, 522, 523, 524, 525, 526, 527, 528, 529, 530, 531, 532, 533)

    if tipoc.id == 10:
        nums = (535, 536, 537, 538, 539, 540, 541, 542, 543, 544, 545, 546, 547, 548, 549, 550, 552, 553, 554, 555, 556, 557,  558, 820, 821, 559, 560, 818, 819, 561, 562, 563, 564, 565, 566,
                567, 568, 569, 570, 571, 572, 573, 574, 575, 576, 577, 578, 579, 580, 581, 582, 583, 584, 585, 586, 587, 588, 589, 590, 591, 851, 852, 853, 854, 592, 593, 594, 595, 596, 597, 598, 599, 600)
    # secue = Secuencia.objects.get(id=2)

    if tipoc.id == 4:
        nums = (129, 130, 131, 132, 133, 134, 135, 136, 137, 138, 139, 140, 141, 142, 143, 144, 145, 146, 147, 148, 149, 150, 151, 152, 153, 154, 155, 156, 157,
                158, 159, 160, 161, 162, 163, 164, 165, 166, 167, 168, 169, 170, 171, 172, 173, 174, 175, 176, 177, 178, 179, 180, 181, 182, 183, 184, 185, 186, 187)

    if tipoc.id == 14:
        nums = (759, 760, 761, 762, 763, 764, 765, 766, 767, 768, 769, 770, 771, 772, 773, 774, 775, 776, 777, 778, 779, 780, 781, 782, 783, 784, 785, 786, 787,
                788, 789, 790, 791, 792, 793, 794, 795, 796, 797, 798, 799, 800, 801, 802, 803, 804, 805, 806, 807, 808, 855, 856, 857, 858, 809, 810, 811, 812, 813, 814, 815, 816, 817)

    if tipoc.id == 5:
        nums = (190, 191, 192, 193, 194, 195, 196, 197, 198, 199, 200, 201, 202, 204, 203, 839, 838,  205, 206, 207, 208, 209, 210, 211, 212, 213, 214, 215, 216, 217,
                218, 219, 220, 221, 222, 223, 224, 225, 226, 227, 228, 229, 230, 231, 232, 233, 234, 235, 236, 237, 238, 239, 850, 847, 848, 849, 240, 241, 242, 243, 244, 245, 246)

    if tipoc.id == 6:
        nums = (249, 250, 251, 252, 253, 254, 255, 256, 257, 258, 960, 961, 962, 963, 259, 260, 261, 262, 263, 841, 840, 264, 265, 266, 267, 268, 269, 270, 271, 272, 273, 274, 275, 276, 277, 278, 279,
                280, 281, 282, 283, 284, 285, 286, 287, 288, 289, 290, 291, 292, 293, 294, 295, 296, 297, 298, 299, 300, 301, 302, 303, 304, 305, 306, 307, 308, 309, 310, 311, 312)

    if tipoc.id == 2:
        nums = (38, 39, 40, 41, 42, 43, 44, 45, 46, 47, 48, 49, 50, 51, 52, 53, 54, 55, 56, 57,
                58, 59, 60, 61, 62, 63, 64, 65, 66, 67, 68, 69, 70, 71, 72, 73, 74, 75, 76, 77, 78, 79)

    if tipoc.id == 3:
        nums = (82, 83, 84, 85, 86, 87, 88, 89, 90, 91, 92, 93, 94, 95, 96, 97, 98, 99, 100, 101, 102, 103, 104, 105,
                106, 107, 108, 109, 110, 111, 112, 113, 114, 115, 116, 117, 118, 119, 120, 121, 122, 123, 124, 125, 126)

    if tipoc.id == 11:
        nums = (604, 605, 606, 607, 608, 609, 610, 611, 612, 613, 614, 615, 616, 617, 618, 619, 620, 621, 622, 623, 624, 625,
                626, 627, 628, 629,  630, 631, 632, 633, 634, 635, 636, 637, 638, 639, 640, 641, 642, 643, 644, 645, 646, 648)

    if tipoc.id == 13:
        nums = (650, 651, 652, 653, 654, 655, 681, 656, 657, 658, 659, 660, 661, 662, 663, 664, 665, 666, 667, 668, 669, 670, 671, 672, 673,
                674, 675, 676, 677, 678, 679, 680)

    if tipoc.id == 15:
        nums = (869, 870, 871, 872, 950, 874, 875, 876, 877, 878, 956, 951, 957, 958, 959, 879, 880, 881, 882, 938, 939, 940, 941, 942, 943, 944, 945, 946, 947, 948, 949, 893, 894, 895, 896, 897, 898, 952, 953, 899, 900, 901, 902, 903, 904, 905, 906, 907, 908, 909, 910, 911, 912, 913, 914, 915, 916, 917, 918, 919, 920, 921, 922, 923, 924, 925, 926, 927, 928, 929, 930, 931, 954, 955, 932, 933, 934, 935, 936)

    if tipoc.id == 16:
        nums = (965, 966, 967, 968, 969, 970, 971, 972, 973, 974, 1035, 1031, 1032, 1033, 1034, 975, 976, 977, 978, 979, 1029, 1030, 980, 981, 982, 983, 984, 985, 986, 987, 988, 989, 990, 991, 992, 993, 994, 995, 996, 997, 998, 999, 1000, 1001, 1002, 1003, 1004, 1005, 1006, 1007, 1008, 1009, 1010, 1011, 1012, 1013, 1014, 1016, 1017, 1018, 1019, 1020, 1021, 1022, 1023, 1024, 1025, 1026, 1027, 1028 )         

    for n in nums:
        # print(n)
        secue = Secuencia.objects.get(id=n)
        if secue.nivel2 == 0:

            secue = Secuencia.objects.get(id=n)
            p005 = document.add_paragraph()
            if secue.identificador != "," and secue.identificador != ".":
                textosecue = secue.identificador + ".- " + secue.textoSecuencia
            else:
                textosecue = secue.textoSecuencia

            textosecue = textosecue.replace(
                "@enCalidadDe1", tipoc.enCalidadDe1)

            if partes.personaParte == 1:
                xcurp = partes.curp[10:11]
                if xcurp == "H":
                    xelolaParte = "EL"
                    xenCalidadDe2 = tipoc.enCalidadDe2
                    textosecue = textosecue.replace(
                        "@enCalidadDe2", xenCalidadDe2)
                else:
                    xelolaParte = "LA"
                    xenCalidadDe2 = tipoc.enCalidadDe2f
                    textosecue = textosecue.replace(
                        "@enCalidadDe2", xenCalidadDe2)

            else:
                xelolaParte = ""
                xenCalidadDe2 = tipoc.enCalidadDe2
                textosecue = textosecue.replace("@enCalidadDe2", xenCalidadDe2)

            if partes.personaParte == 1:
                if partes.fecha_ingreso:
                    # print(secue.id)
                    textosecue = textosecue.replace(
                        "@fechaingreso", partes.fecha_ingreso.strftime("%d de %B de %Y"))

                else:
                    textosecue = textosecue.replace(
                        "@fechaingreso", "**********")

                if puesto.funcionesPuesto:
                    #     # print(secue.id)

                    textosecue = textosecue.replace(
                        "@actividadesPuesto", puesto.funcionesPuesto)

                else:
                    textosecue = textosecue.replace(
                        "@actividadesPuesto", "**********")

                if partes.actividadesParte:
                    # print(secue.id)
                    textosecue = textosecue.replace(
                        "@actividadesContrato", partes.actividadesParte)
                else:
                    textosecue = textosecue.replace(
                        "@actividadesContrato", "")

                if partes.titulo_profParte:
                    textosecue = textosecue.replace(
                        "@titulo_profParte", partes.titulo_profParte)
                else:
                    textosecue = textosecue.replace(
                        "@titulo_profParte", "**********")

                if partes.clavePuesto:
                    # print(secue.id)
                    textosecue = textosecue.replace(
                        "@clavePuesto", puesto.nombrePuesto)
                else:
                    textosecue = textosecue.replace(
                        "@clavePuesto", "**********")

            textosecue = textosecue.replace(
                "@totalhorasContrato", str(contratos.totalhorasContrato))
            textosecue = textosecue.replace("@importeContrato", currency)
            textosecue = textosecue.replace("@letras", "(" + letras + ")")
            # print(contratos.datecontrato_ini)
            # print(contratos.datecontrato_fin)
            textosecue = textosecue.replace(
                "@datecontrato_ini", contratos.datecontrato_ini.strftime("%d de %B de %Y"))
            textosecue = textosecue.replace(
                "@datecontratox", datecontratox.strftime("%d de %B de %Y"))

            if partes.beneficiario1:
                textosecue = textosecue.replace(
                    "@beneficiario1", partes.beneficiario1.ljust(50)[:50])
            else:
                textosecue = textosecue.replace(
                    "@beneficiario1", "**********")

            if partes.beneficiario2:
                textosecue = textosecue.replace(
                    "@beneficiario2", partes.beneficiario2.ljust(50)[:50])
            else:
                textosecue = textosecue.replace(
                    "@beneficiario2", "")    

            if partes.beneficiario3:
                textosecue = textosecue.replace(
                    "@beneficiario3", partes.beneficiario3.ljust(50)[:50])
            else:
                textosecue = textosecue.replace(
                    "@beneficiario3", "")

            if partes.parentesco1:
                textosecue = textosecue.replace(
                    "@parentesco1", partes.parentesco1)
            else:
                textosecue = textosecue.replace(
                    "@parentesco1", "**********")

            if partes.parentesco2:
                textosecue = textosecue.replace(
                    "@parentesco2", partes.parentesco2)
            else:
                textosecue = textosecue.replace(
                    "@parentesco2", "")

            if partes.parentesco3:
                textosecue = textosecue.replace(
                    "@parentesco3", partes.parentesco3)
            else:
                textosecue = textosecue.replace(
                    "@parentesco3", "")    

            if partes.porcentaje1:
                textosecue = textosecue.replace(
                    "@porcentaje1", partes.porcentaje1)
            else:
                textosecue = textosecue.replace(
                    "@porcentaje1", "**********")

            if partes.porcentaje2:
                textosecue = textosecue.replace(
                    "@porcentaje2", partes.porcentaje2)
            else:
                textosecue = textosecue.replace(
                    "@porcentaje2", "")

            if partes.porcentaje3:
                textosecue = textosecue.replace(
                    "@porcentaje3", partes.porcentaje3)
            else:
                textosecue = textosecue.replace(
                    "@porcentaje3", "")    

            textosecue = textosecue.replace(
                "@area", departamento.nombreDepartamento)

            textosecue = textosecue.replace(
                "@direccioncampus", campus.direccionCampus)

            if contratos.datecontrato_fin:
                textosecue = textosecue.replace(
                    "@datecontrato_fin", contratos.datecontrato_fin.strftime("%d de %B de %Y"))
            textosecue = textosecue.replace(
                "@datecontrato", contratos.datecontrato.strftime("%d de %B de %Y"))

            textosecue = textosecue.replace(
                "@npContrato", str(contratos.npContrato))
            textosecue = textosecue.replace("@imppContrato", currency2)
            textosecue = textosecue.replace(
                "@pagolet", "(" + pagolet + ")")

            textosecue = textosecue.replace(
                "January", "Enero")
            textosecue = textosecue.replace(
                "February", "Febrero")
            textosecue = textosecue.replace(
                "March", "Marzo")
            textosecue = textosecue.replace(
                "April", "Abril")
            textosecue = textosecue.replace(
                "May", "Mayo")
            textosecue = textosecue.replace(
                "June", "Junio")
            textosecue = textosecue.replace(
                "July", "Julio")
            textosecue = textosecue.replace(
                "August", "Agosto")
            textosecue = textosecue.replace(
                "September", "Septiembre")
            textosecue = textosecue.replace(
                "October", "Octubre")
            textosecue = textosecue.replace(
                "November", "Noviembre")
            textosecue = textosecue.replace(
                "December", "Diciembre")

            if xcurp == "M":
                textosecue = textosecue.replace(
                    "CATEDRÁTICO UNIVERSITARIO Y ASESOR", "CATEDRÁTICA UNIVERSITARIA Y ASESORA")
                textosecue = textosecue.replace("este último", "esta última")
            if secue.identificador == ".":
                p005.add_run(textosecue, style='CommentsStyle').bold = False
            else:
                p005.add_run(textosecue, style='CommentsStyle').bold = True
            paragraph_format = p005.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.left_indent = Inches(0.0)
            # print(len(textosecue))
            # print(textosecue)

            if len(textosecue) > 100 and len(textosecue) < 1500:
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            else:
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        else:

            secue = Secuencia.objects.get(id=n)
            p006 = document.add_paragraph()
            textosecue = secue.identificador + ".- " + secue.textoSecuencia

            if partes.datos_actaconstParte:
                textosecue = textosecue.replace(
                    "@datos_actaconstParte", partes.datos_actaconstParte)

            if partes.curp:
                textosecue = textosecue.replace("@curp", partes.curp)

            if partes.titulo_profParte:
                textosecue = textosecue.replace(
                    "@titulo_profParte", partes.titulo_profParte)
            else:
                textosecue = textosecue.replace(
                    "@titulo_profParte", "**********")

            textosecue = textosecue.replace(
                "@enCalidadDe1", tipoc.enCalidadDe1)
            textosecue = textosecue.replace("@enCalidadDe2", xenCalidadDe2)

            if partes.universidadParte:
                textosecue = textosecue.replace(
                    "@universidadParte", partes.universidadParte)
            else:
                textosecue = textosecue.replace(
                    "@universidadParte", "**********")

            if partes.cedula_profParte:
                textosecue = textosecue.replace(
                    "@cedula_profParte", partes.cedula_profParte)
            else:
                textosecue = textosecue.replace(
                    "@cedula_profParte", "**********")

            if ciclo.ciclo_actual:
                textosecue = textosecue.replace(
                    "@CicloParte", ciclo.descripcionCiclo)
            else:
                textosecue = textosecue.replace("@CicloParte", "**********")

            if partes.rfc:
                textosecue = textosecue.replace("@rfc", partes.rfc)
            else:
                textosecue = textosecue.replace("@rfc", "**********")

            if regimen.nombreRegimen:
                textosecue = textosecue.replace(
                    "@regfiscalParte", regimen.nombreRegimen)
            else:
                textosecue = textosecue.replace(
                    "@regfiscalParte", "**********")

            if partes.domicilioParte:
                textosecue = textosecue.replace(
                    "@domicilioParte", partes.domicilioParte)
            else:
                textosecue = textosecue.replace(
                    "@domicilioParte", "**********")

            if contratos.marcaVehiculo:
                textosecue = textosecue.replace(
                    "@marcaVehiculo", contratos.marcaVehiculo)
            else:
                textosecue = textosecue.replace(
                    "@marcaVehiculo", "**********")

            if contratos.empresa:
                textosecue = textosecue.replace(
                    "@empresa", contratos.empresa)
            else:
                textosecue = textosecue.replace(
                    "@empresa", "**********")

            if contratos.tipo_sociedad:
                textosecue = textosecue.replace(
                    "@tipo_sociedad", contratos.tipo_sociedad)
            else:
                textosecue = textosecue.replace(
                    "@tipo_sociedad", "**********")

            if contratos.objeto_social:
                textosecue = textosecue.replace(
                    "@objeto_social", contratos.objeto_social)
            else:
                textosecue = textosecue.replace(
                    "@objeto_social", "**********")    

            if contratos.testimonio:
                textosecue = textosecue.replace(
                    "@testimonio", contratos.testimonio)
            else:
                textosecue = textosecue.replace(
                    "@testimonio", "**********")

            if contratos.clausula:
                textosecue = textosecue.replace(
                    "@clausula", contratos.clausula)
            else:
                textosecue = textosecue.replace(
                    "@clausula", "**********")

            if contratos.domicilio_sociedad:
                textosecue = textosecue.replace(
                    "@domicilioSociedad", contratos.domicilio_sociedad)
            else:
                textosecue = textosecue.replace(
                    "@domicilioSociedad", "**********")   

            if contratos.rfc_sociedad:
                textosecue = textosecue.replace(
                    "@r_fcSociedad", contratos.rfc_sociedad)
            else:
                textosecue = textosecue.replace(
                    "@r_fcSociedad", "**********")

            if contratos.fecha_constitucion:
                textosecue = textosecue.replace(
                    "@fechaconstitucion", contratos.fecha_constitucion)
            else:
                textosecue = textosecue.replace(
                    "@fechaconstitucion", "**********")

            if contratos.registro_mercantil:
                textosecue = textosecue.replace(
                    "@registromercantil", contratos.registro_mercantil)
            else:
                textosecue = textosecue.replace(
                    "@registromercantil", "**********")        

            if contratos.parte2:
                textosecue = textosecue.replace(
                    "@nombreParte2", xnombreParte)
            else:
                textosecue = textosecue.replace(
                    "@nombreParte2", "**********")

            if contratos.modeloVehiculo:
                textosecue = textosecue.replace(
                    "@modeloVehiculo", contratos.modeloVehiculo)
            else:
                textosecue = textosecue.replace(
                    "@modeloVehiculo", "**********")

            if contratos.tipoVehiculo:
                textosecue = textosecue.replace(
                    "@tipoVehiculo", contratos.tipoVehiculo)
            else:
                textosecue = textosecue.replace(
                    "@tipoVehiculo", "**********")

            if contratos.motorVehiculo:
                textosecue = textosecue.replace(
                    "@motorVehiculo", contratos.motorVehiculo)
            else:
                textosecue = textosecue.replace(
                    "@motorVehiculo", "**********")

            if contratos.serieVehiculo:
                textosecue = textosecue.replace(
                    "@serieVehiculo", contratos.serieVehiculo)
            else:
                textosecue = textosecue.replace(
                    "@serieVehiculo", "**********")

            if contratos.placasVehiculo:
                textosecue = textosecue.replace(
                    "@placasVehiculo", contratos.placasVehiculo)
            else:
                textosecue = textosecue.replace(
                    "@placasVehiculo", "**********")

            if contratos.facturaVehiculo:
                textosecue = textosecue.replace(
                    "@facturaVehiculo", contratos.facturaVehiculo)
            else:
                textosecue = textosecue.replace(
                    "@facturaVehiculo", "**********")

            if contratos.fechaFactura:
                textosecue = textosecue.replace(
                    "@fechaFactura", contratos.fechaFactura.strftime("%d de %B de %Y"))
            else:
                textosecue = textosecue.replace(
                    "@fechaFactura", "**********")

            if contratos.expedidaFactura:
                textosecue = textosecue.replace(
                    "@expedidaFactura", contratos.expedidaFactura)
            else:
                textosecue = textosecue.replace(
                    "@expedidaFactura", "**********")

            if contratos.tarjetaVehiculo:
                textosecue = textosecue.replace(
                    "@tarjetaVehiculo", contratos.tarjetaVehiculo)
            else:
                textosecue = textosecue.replace(
                    "@tarjetaVehiculo", "**********")

            if contratos.fechaTarjeta:
                textosecue = textosecue.replace(
                    "@fechaTarjeta", contratos.fechaTarjeta.strftime("%d de %B de %Y"))
            else:
                textosecue = textosecue.replace(
                    "@fechaTarjeta", "**********")

            if contratos.expedidaFactura:
                textosecue = textosecue.replace(
                    "@expedidaTarjeta", contratos.expedidaTarjeta)
            else:
                textosecue = textosecue.replace(
                    "@expedidaTarjeta", "**********")

            if contratos.polizaVehiculo:
                textosecue = textosecue.replace(
                    "@polizaVehiculo", contratos.polizaVehiculo)
            else:
                textosecue = textosecue.replace(
                    "@polizaVehiculo", "**********")

            if contratos.fechaPoliza:
                textosecue = textosecue.replace(
                    "@fechaPoliza", contratos.fechaPoliza.strftime("%d de %B de %Y"))
            else:
                textosecue = textosecue.replace(
                    "@fechaPoliza", "**********")

            if partes.clavePuesto:
                # print(secue.id)
                textosecue = textosecue.replace(
                        "@clavePuesto", puesto.nombrePuesto)
            else:
                textosecue = textosecue.replace(
                        "@clavePuesto", "**********")    

            if contratos.expedidaPoliza:
                textosecue = textosecue.replace(
                    "@expedidaPoliza", contratos.expedidaPoliza)
            else:
                textosecue = textosecue.replace(
                    "@expedidaPoliza", "**********")

            if partes.nacionalidadParte:
                textosecue = textosecue.replace(
                    "@nacionalidadParte", partes.nacionalidadParte)
            else:
                textosecue = textosecue.replace(
                    "@nacionalidadParte", "**********")

            ano = partes.rfc[4:6]
            mes = partes.rfc[6:8]
            dia = partes.rfc[8:10]

            # Convertir el año a un entero
            xano = int(ano)

            # Si el año es menor que 50, asumimos que es 2000 o posterior
            if xano < 40:
                xano += 2000
            else:
                xano += 1900

            if partes.personaParte == 1:
                fecha_nacimiento = date(xano, int(mes), int(dia))
                edad = calcular_edad_anos(fecha_nacimiento)
                # print(f'la edad es {edad} años')

                if edad:
                    textosecue = textosecue.replace(
                        "@edadParte", " "+str(edad))
                else:
                    textosecue = textosecue.replace("@edadParte", "**********")

                sexo = partes.curp[10:11]

                if sexo == "H":
                    textosecue = textosecue.replace("@sexoParte", "MASCULINO")
                else:
                    if sexo == "M":
                        textosecue = textosecue.replace(
                            "@sexoParte", "FEMENINO")
                    else:
                        textosecue = textosecue.replace(
                            "@sexoParte", "**********")

                if partes.estadocivilParte:
                    textosecue = textosecue.replace(
                        "@estadocivilParte", partes.estadocivilParte)
                else:
                    textosecue = textosecue.replace(
                        "@estadocivilParte", "**********")

            textosecue = textosecue.replace(
                "@datecontrato_ini", contratos.datecontrato_ini.strftime("%d de %B de %Y"))
            if contratos.datecontrato_fin:
                textosecue = textosecue.replace(
                    "@datecontrato_fin", contratos.datecontrato_fin.strftime("%d de %B de %Y"))
            textosecue = textosecue.replace(
                "@datecontrato", contratos.datecontrato.strftime("%d de %B de %Y"))

            textosecue = textosecue.replace(
                "@idrep_legalParte", replegal.tituloParte + " " + replegal.nombreParte)
            textosecue = textosecue.replace("@RPImssParteC", "8401667310-9 ")
            textosecue = textosecue.replace(
                "@enCalidadDe1", contratos.enCalidadDe1)
            textosecue = textosecue.replace(
                "@enCalidadDe2", contratos.enCalidadDe2)
            textosecue = textosecue.replace(
                "@domicilioPatron", patron.domicilioParte)
            textosecue = textosecue.replace("@importeContrato", currency)
            textosecue = textosecue.replace("@letras", "(" + letras + ")")
            textosecue = textosecue.replace(
                "@npContrato", str(contratos.npContrato))

            textosecue = textosecue.replace(
                "@totalhorasContrato", str(contratos.totalhorasContrato))
            if currency2:
                textosecue = textosecue.replace("@imppContrato", currency2)
                textosecue = textosecue.replace(
                    "@pagolet", "(" + pagolet + ")")
            if xcurp == "M":
                textosecue = textosecue.replace("mexicano", "mexicana")
                textosecue = textosecue.replace(
                    "un profesionista", "una profesionista")
                textosecue = textosecue.replace("inscrito", "inscrita")

            if puesto.funcionesPuesto:
                # print(secue.id)
                textosecue = textosecue.replace(
               "@actividadesPuesto", puesto.funcionesPuesto)
            else:
                textosecue = textosecue.replace(
               "@actividadesPuesto", "***********")

            p006.add_run(textosecue, style='CommentsStyle').bold = False
            paragraph_format = p006.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(3)
            paragraph_format.left_indent = Inches(0.4)
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if partes.personaParte == 1:
                xcurp = partes.curp[10:11]

    dtable = document.add_table(rows=3, cols=2)
    dtable.style = "TableNormal"
    dtable.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    dtable.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    dtable.cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    dtable.cell(1, 0).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    dtable.cell(2, 1).vertical_alignment = WD_ALIGN_VERTICAL.TOP
    dtable.cell(2, 0).vertical_alignment = WD_ALIGN_VERTICAL.TOP

    dtab_cells = dtable.rows[0].cells
    dt1 = dtab_cells[0].text = ''
    dt1 = dtab_cells[0].paragraphs[0].add_run(
        tipoc.enCalidadDe1).font.size = Pt(11)
    dt1 = dtab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt1 = dtab_cells[1].text = ''
    dt1 = dtab_cells[1].paragraphs[0].add_run(xenCalidadDe2).font.size = Pt(11)
    dt1 = dtab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dtab_cells = dtable.rows[1].cells
    dtable.rows[1].height_rule = WD_ROW_HEIGHT.EXACTLY
    dtable.rows[1].height = 342900
    dt1 = dtab_cells[0].text = ''
    dt1 = dtab_cells[0].paragraphs[0].add_run(
        "___________________________________").font.size = Pt(11)
    dt1 = dtab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt1 = dtab_cells[1].text = ''
    dt1 = dtab_cells[1].paragraphs[0].add_run(
        "___________________________________").font.size = Pt(11)
    dt1 = dtab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dtab_cells = dtable.rows[2].cells
    dt2 = dtab_cells[0].text = ''
    dt2 = dtab_cells[0].paragraphs[0].add_run(
        replegal.tituloParte + " " + replegal.nombreParte).font.size = Pt(11)
    dt2 = dtab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt2 = dtab_cells[1].text = ''

    if partes.tituloParte:
        dt2 = dtab_cells[1].paragraphs[0].add_run(
            partes.tituloParte + " " + partes.nombreParte).font.size = Pt(11)
    else:
        if partes.personaParte == 1:
            dt2 = dtab_cells[1].paragraphs[0].add_run(
                "**********" + " " + partes.nombreParte).font.size = Pt(11)
        else:
            dt2 = dtab_cells[1].paragraphs[0].add_run(
                partes.nombreParte).font.size = Pt(11)

    dt2 = dtab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dtable = document.add_table(rows=1, cols=2)

    dtable1 = document.add_table(rows=3, cols=2)
    dtable1.style = "TableNormal"
    dtable1.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    dtable1.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    dtable1.cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    dtable1.cell(1, 0).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    dtable1.cell(2, 1).vertical_alignment = WD_ALIGN_VERTICAL.TOP
    dtable1.cell(2, 0).vertical_alignment = WD_ALIGN_VERTICAL.TOP

    dtab_cells = dtable1.rows[0].cells
    dt3 = dtab_cells[0].text = ''
    dt3 = dtab_cells[0].paragraphs[0].add_run("TESTIGO").font.size = Pt(11)
    dt3 = dtab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt3 = dtab_cells[1].text = ''
    dt3 = dtab_cells[1].paragraphs[0].add_run("TESTIGO").font.size = Pt(11)
    dt3 = dtab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dtab_cells = dtable1.rows[1].cells
    dtable1.rows[1].height_rule = WD_ROW_HEIGHT.EXACTLY
    dtable1.rows[1].height = 342900
    dt4 = dtab_cells[0].text = ''
    dt4 = dtab_cells[0].paragraphs[0].add_run(
        "___________________________________").font.size = Pt(11)
    dt4 = dtab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt4 = dtab_cells[1].text = ''
    dt4 = dtab_cells[1].paragraphs[0].add_run(
        "___________________________________").font.size = Pt(11)
    dt4 = dtab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dtab_cells = dtable1.rows[2].cells
    dt5 = dtab_cells[0].text = ''
    dt5 = dtab_cells[0].paragraphs[0].add_run(
        contratos.testigoContrato1).font.size = Pt(11)
    dt5 = dtab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt5 = dtab_cells[1].text = ''
    dt5 = dtab_cells[1].paragraphs[0].add_run(
        contratos.testigoContrato2).font.size = Pt(11)
    dt5 = dtab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()
    document.add_paragraph()
    xtable = document.add_table(rows=1, cols=2)
    xtable.style = "TableGrid"

    xtab_cells = xtable.rows[0].cells
    xt1 = xtab_cells[1].text = ''

    if partes.tituloParte:
        xt1 = xtab_cells[1].paragraphs[0].add_run(
            partes.tituloParte + " " + partes.nombreParte).font.size = Pt(8)
    else:
        if partes.personaParte == 1:
            xt1 = xtab_cells[1].paragraphs[0].add_run(
                "**********" + " " + partes.nombreParte).font.size = Pt(8)
        else:
            xt1 = xtab_cells[1].paragraphs[0].add_run(
                partes.nombreParte).font.size = Pt(8)

    xt1 = xtab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    xt1 = xtab_cells[0].text = ''
    xt1 = xtab_cells[0].paragraphs[0].add_run(
        "DATOS PARA CONTROL EN EL ÁREA DE RECURSOS HUMANOS").font.size = Pt(8)
    xt1 = xtab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytable = document.add_table(rows=2, cols=5)
    ytable.style = "TableGrid"

    ytab_cells = ytable.rows[0].cells
    yt1 = ytab_cells[0].text = ''
    yt1 = ytab_cells[0].paragraphs[0].add_run(
        "Total Contrato en Horas").font.size = Pt(8)
    yt1 = ytab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[0].cells
    yt1 = ytab_cells[1].text = ''
    yt1 = ytab_cells[1].paragraphs[0].add_run(
        "Total Contrato en $").font.size = Pt(8)
    yt1 = ytab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[0].cells
    yt1 = ytab_cells[2].text = ''
    yt1 = ytab_cells[2].paragraphs[0].add_run(
        "Clave del Depto.").font.size = Pt(8)
    yt1 = ytab_cells[2].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[0].cells
    yt1 = ytab_cells[3].text = ''
    yt1 = ytab_cells[3].paragraphs[0].add_run(
        "Ingreso o Reingreso").font.size = Pt(8)
    yt1 = ytab_cells[3].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[0].cells
    yt1 = ytab_cells[4].text = ''
    yt1 = ytab_cells[4].paragraphs[0].add_run("Versión").font.size = Pt(8)
    yt1 = ytab_cells[4].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[1].cells
    yt2 = ytab_cells[0].text = ''
    yt2 = ytab_cells[0].paragraphs[0].add_run(
        str(contratos.totalhorasContrato)).font.size = Pt(8)
    yt2 = ytab_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[1].cells
    yt2 = ytab_cells[1].text = ''
    yt2 = ytab_cells[1].paragraphs[0].add_run(currency).font.size = Pt(8)
    yt2 = ytab_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[1].cells
    yt2 = ytab_cells[2].text = ''
    yt2 = ytab_cells[2].paragraphs[0].add_run(
        str(partes.claveDepartamento)).font.size = Pt(8)
    yt2 = ytab_cells[2].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[1].cells
    yt2 = ytab_cells[3].text = ''
    yt2 = ytab_cells[3].paragraphs[0].add_run(" ").font.size = Pt(8)
    yt2 = ytab_cells[3].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ytab_cells = ytable.rows[1].cells
    yt2 = ytab_cells[4].text = ''
    yt2 = ytab_cells[4].paragraphs[0].add_run(
        "EMODELO 17 v5 enero 2019").font.size = Pt(8)
    yt2 = ytab_cells[4].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Added new section for assigning different footer on each page.
    new_section = document.sections[0].footer
    sectPr = new_section._sectPr

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), 'decimal')
    pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)

    # Get footer-area of the recent section in document
    new_footer = document.sections[0].footer
    new_footer.is_linked_to_previous = False

    footer_para = new_footer.paragraphs[0]
    run_footer = footer_para.add_run("Pág- ").font.size = Pt(11)
    run_footer = footer_para.add_run(pgNumType)
    _add_number_range(run_footer)
    font = run_footer.font
    font.name = 'Arial'
    font.size = Pt(11)
    run_footer = new_footer.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Print the user's name
    # document_elements_heading = document.add_heading(document_data_full_name, 0)
    # document_elements_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Print biography and careerpath

    # Add empty paragraph
    # document.add_paragraph()

    # Sincerely and name
    # document.add_paragraph("Atentamente,\n" + document_data_full_name)
    document.add_page_break()

    # Add an empty paragraph to advance one line
    document.add_paragraph()

    # Add the text "Secuencia de un registro de Secuencia" after the page break
    paragraph = document.add_paragraph("Aviso de Privacidad Integral del expediente de personal y registro de asistencia")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set the font to Arial and size to 11
    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = True

    # Add the remaining text
    paragraph = document.add_paragraph("La empresa “ESCUELA MODELO S.C.P.”, con domicilio en calle 56 A no 444 x 29 A Col Centro, Mérida, Yucatán, con código postal 97000, es responsable del tratamiento de los datos personales, que nos proporcione toda persona que trabaje, colabore, preste sus servicios personales y/o servicios especializados para la empresa, los cuales serán protegidos conforme a lo dispuesto por la Ley Federal de Protección de Datos Personales en Posesión de los Particulares, y demás normatividad que resulte aplicable.")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = False

    paragraph = document.add_paragraph("Finalidades del tratamiento")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set the font to Arial and size to 11
    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = True

    paragraph = document.add_paragraph("Los datos personales que recabamos de usted, persona que trabaje, colabore, preste sus servicios personales para la empresa, los utilizaremos para las siguientes finalidades: realizar los trámites de contratación, nombramiento e identificación de personal; administrar y dispersar la nómina; cumplir con las obligaciones patronales; otorgamiento de prestaciones y movimientos de personal, cumplimiento de obligaciones de transparencia comunes establecidos en la Ley Federal de Protección de Datos Personales en Posesión de los Particulares, transferencia a terceros en cumplimiento a atribuciones legales, registro de asistencia electrónica; así mismo se comunica que no se efectuarán tratamientos adicionales.")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = False

    paragraph = document.add_paragraph("Datos personales recabados")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set the font to Arial and size to 11
    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = True

    paragraph = document.add_paragraph("Para las finalidades antes señaladas se solicitarán los siguientes datos personales: .")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = False

    datos_table = document.add_table(rows=1, cols=2)
    datos_table.style = 'Table Grid'
    header_cells = datos_table.rows[0].cells
    header_cells[0].text = "Categoría"
    header_cells[1].text = "Tipo de datos personales"

    rows = [
        ("Datos identificativos", "Nombre, Domicilio, Teléfono particular, Teléfono celular, Estado Civil, Firma, "
                              "Registro Federal de Contribuyentes, Código Postal para efectos fiscales, "
                              "Régimen de contribución en términos LISR, Clave Única de Registro de la Población, "
                              "Nombre de familiares, dependientes y beneficiarios, Fecha de nacimiento, Lugar de nacimiento, "
                              "Fotografías físicas, Edad."),
    ("Datos electrónicos", "Correo electrónico, Fotografías en el centro de trabajo, Videograbaciones en el centro de trabajo."),
    ("Datos académicos", "Títulos, Certificados, Reconocimientos, Constancias, Diplomas, Cédula Profesional."),
    ("Datos laborales", "Nombramiento, Referencias personales y laborales, Número de seguro social."),
    ("Datos patrimoniales", "Seguros, Número de cuenta y/o Clave Bancaria Estandarizada, Información Fiscal/ Constancia "
                            "de Situación Fiscal, Descuentos por orden judicial, Créditos, Ingresos."),
    ("Datos biométricos", "Rostro, Huella dactilar.")
    ]
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
    
    for categoria, tipo in rows:
        row_cells = datos_table.add_row().cells
        row_cells[0].text = categoria
        # row_cells[1].text = tipo
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        tipos = tipo.split(',')
        for t in tipos:
            p = row_cells[1].add_paragraph()
            run = p.add_run(t.strip())
            run.bold = False
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            p.style = 'ListBullet'
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = Pt(11)

        if len(rows) == 1:
            for cell in row_cells:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls("w")))
                )
        else:
            for cell in row_cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = False
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    
    
    document.add_paragraph()
    paragraph = document.add_paragraph("Transferencia de datos personales ")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set the font to Arial and size to 11
    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = True

    paragraph = document.add_paragraph("Le informamos que sus datos personales son compartidos con: .")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = False
    
    transferencia_table = document.add_table(rows=1, cols=3)
    transferencia_table.style = 'Table Grid'
    transferencia_table.columns[0].width = Inches(4.5) 
    header_cells = transferencia_table.rows[0].cells
    header_cells[0].text = "Destinatario de los datos personales"
    header_cells[1].text = "País"
    header_cells[2].text = "Finalidad"
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(11)
    
    transferencia_rows = [
        ("Servicio de Administración Tributaria y Unidad de Inteligencia Financiera de la Secretaría de Hacienda y Crédito Público", 
         "México", "Pago de impuestos, actividades vulnerables, registros fiscales, controles volumétricos."),
        ("Instituto Mexicano del Seguro Social", 
         "México", "Pago de cuotas obrero-patronales."),
        ("Instituto del Fondo Nacional para la Vivienda de los Trabajadores", 
         "México", "Pago de Aportaciones y Amortizaciones."),
        ("Instituto del Fondo Nacional para el Consumo de los Trabajadores", 
         "México", "Pago de amortizaciones de créditos."),
        ("Secretaría de Finanzas del Estado de Yucatán o Quintana Roo", 
         "México", "Trámites financieros y nómina."),
        ("BBVA México, S.A., Institución de Banca Múltiple, Grupo Financiero BBVA México y/o HSBC México, S.A., Institución de Banca Múltiple, Grupo Financiero HSBC", 
         "México", "Dispersión de nómina."),
        ("Autoridades jurisdiccionales estatales o federales", 
         "México", "Cumplimiento de mandamiento judicial fundado y motivado.")
    ]

    for destinatario, pais, finalidad in transferencia_rows:
        row_cells = transferencia_table.add_row().cells
        row_cells[0].text = destinatario
        row_cells[1].text = pais
        row_cells[2].text = finalidad
    
        # Apply font properties to each cell
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = False
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                        
    document.add_paragraph()
    
    
    paragraph = document.add_paragraph("Le informamos que, para las transferencias indicadas, en el rubro Finalidad requerimos obtener su consentimiento. Si usted no manifiesta su negativa para dichas transferencias, entenderemos que nos lo ha otorgado. ")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set the font to Arial and size to 11
    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = False               
    
    document.add_page_break()
       
   
    paragraph = document.add_paragraph("Datos Personales Sensibles. ")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set the font to Arial and size to 11
    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = True
    
    paragraph = document.add_paragraph("Con fundamento legal, en la fracción VI, Artículo 3 y Artículo 9 de la Ley Federal de Protección de Datos Personales en Posesión de los Particulares, la persona que trabaja, colabora, presta sus servicios personales para la empresa, señala que: SI / NO tiene datos personales que afecten a la esfera más íntima de su persona, o cuya utilización indebida pueda dar origen a discriminación, en términos de los artículos 2 y 3, de la Ley Federal del Trabajo, o conlleve un riesgo grave para éste. En particular, se consideran sensibles aquellos que puedan revelar aspectos como origen racial o étnico, estado de salud presente y futuro, información genética, creencias religiosas, filosóficas y morales, afiliación sindical, opiniones políticas, preferencia sexual, antecedentes penales. Estos datos, son de uso exclusivo e interno de la empresa, con base al Protocolo para prevenir la discriminación, así para dar cumplimiento a nuestra Política de prevención de riesgos psicosociales.  Por lo que SI / NO requisitará el ANEXO DATOS PERSONALES SENSIBLES, para declarar los posibles DATOS PERSONALES SENSIBLES.")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = False
    
    paragraph = document.add_paragraph("Derechos ARCO ")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set the font to Arial and size to 11
    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = True
    
    paragraph = document.add_paragraph("Usted, persona que trabaja, colabora, presta sus servicios personales para la empresa, tiene derecho a conocer qué datos personales tenemos de usted, para qué los utilizamos y las condiciones del uso que les damos (Acceso). Asimismo, es su derecho solicitar la corrección de su información personal en caso de que esté desactualizada, sea inexacta o incompleta (Rectificación); que la eliminemos de nuestros registros o bases de datos cuando considere que la misma no está siendo utilizada conforme a los principios, deberes y obligaciones previstas en la normativa (Cancelación); así como oponerse al uso de sus datos personales para fines específicos (Oposición). Estos derechos se conocen como derechos ARCO. " )
    
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = False
    
    paragraph = document.add_paragraph("Para el ejercicio de cualquiera de los derechos ARCO, usted deberá presentar la solicitud POR ESCRITO respectiva en privacidadmodelo@modelo.edu.mx " )
    
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = False
    
    paragraph = document.add_paragraph("Para conocer el procedimiento y requisitos para el ejercicio de los derechos ARCO, usted podrá ponerse en contacto con nuestro Departamento de Privacidad, que dará trámite a las solicitudes para el ejercicio de estos derechos, y atenderá cualquier duda que pudiera tener respecto al tratamiento de su información. Los datos de contacto del Departamento de Privacidad son los siguientes privacidadmodelo@modelo.edu.mx" )
    
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = False
    
    
    paragraph = document.add_paragraph("Cambios al Aviso de Privacidad ")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set the font to Arial and size to 11
    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = True
    
    paragraph = document.add_paragraph("En caso de realizar alguna modificación al Aviso de Privacidad, se le hará de su conocimiento vía correo electrónico o a través del portal de la empresa: “ESCUELA MODELO S.C.P.”," )
    
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = False
    
    fecha_formateada = format_date(contratos.datecontrato_ini, format='d ' 'MMMM ' 'YYYY', locale='es_ES')
    
    paragraph = document.add_paragraph("El presente Aviso de privacidad, se realizó y se firma el " + fecha_formateada + " bajo protesta de decir verdad, que los datos autorizados y manifestados son ciertos.")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = False
    
    paragraph = document.add_paragraph("__________________________")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar "Nombre completo"
    paragraph = document.add_paragraph()
    run = paragraph.add_run(partes.nombreParte)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar "Doy mi consentimiento total y absoluto para la protección de mis datos"
    
    paragraph = document.add_paragraph("Doy mi consentimiento total y absoluto para la protección de mis datos ")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set the font to Arial and size to 11
    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = False


    # from docx import Document
    # from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
    
    # # Crea el documento
    # doc = Document()
    
    if tipoc.id == 1:# # Estilos globales
        document.add_page_break()
    
        # Add an empty paragraph to advance one line
        # document.add_paragraph() # style = doc.styles['Normal']
        # style.font.name = 'Times New Roman'
        # style.font.size = Pt(12)
        # style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # run = paragraph.runs[0]
        # run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.bold = False
    
        # document.add_paragraph() 
        # p.paragraph_format.space_after = Pt(0)
        # p.paragraph_format.space_before = Pt(0)
    
        fecha_formateada = format_date(contratos.datecontrato_ini, format='d ' 'MMMM ' 'YYYY', locale='es_ES')
        
        paragraph = document.add_paragraph("Mérida, Yucatán a " + fecha_formateada + ".")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.bold = True 
        
        # Espacio
        document.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        # Encabezado de la institución
        for line in [
            "Escuela MODELO S.C.P.",
            "Departamento de Contabilidad y Pagos.",
            "Paseo de Montejo, calle 56-A núm. 444",
            "CP 97000, Mérida, Yucatán."
        ]:
            p = document.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.bold = True
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
        # Espacio
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        # ASUNTO
        p = document.add_paragraph()
        run = p.add_run("ASUNTO: ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run("Solicitud de retención de I.S.R. opcional.")
        run.bold = True
        run.font.size = Pt(11)
        
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        
        nombre_doc = (partes.nombreParte or "").strip()
        domicilio_doc = (partes.domicilioParte or "").strip()
        
        p = document.add_paragraph()
        run = p.add_run("El suscrito, ")
        run.bold = False
        run.font.size = Pt(11)
    
        run = p.add_run(nombre_doc)
        run.bold = True
        run.font.size = Pt(11)
    
        run = p.add_run(" con domicilio fiscal para oír y recibir toda clase de notificaciones y documentos, el ubicado en " +     domicilio_doc + ", de esta ciudad, ante ustedes por mi propio derecho comparezco y expongo lo siguiente:")
        run.bold = False
        run.font.size = Pt(11)
        # p = document.add_paragraph()
        # run = p.add_run(texto1)
        # run.bold = False
        # run.font.size = Pt(11)
        # Segundo párrafo
        texto2 = (
            "Que soy un profesional independiente cuyos ingresos preponderantes del ejercicio fiscal anterior "
            "no provienen de la contribuyente EMO100510EW5 Escuela Modelo S.C.P."
        )
        p = document.add_paragraph()
        run = p.add_run(texto2)
        run.bold = True
        run.font.size = Pt(11)
        # Tercer párrafo
        texto3 = (
            "Que ante su invitación de integrarme a su plantilla laboral como asalariado, manifiesto mi total "
            "y rotunda negativa por las razones que a continuación expongo."
        )
        p = document.add_paragraph()
        run = p.add_run(texto3)
        run.bold = False
        run.font.size = Pt(11)
    
        # Lista numerada
        items = [
            "Que deseo mantener mi independencia profesional, para prestar mis servicios con cualquier otro cliente en los horarios y actividades que a mi interés convengan.",
            "Que actualmente estoy pensionado por el ISSSTE y por el IMSS y gozo de las prestaciones médicas y sociales que ambas instituciones me otorgan."
        ]
        for idx, text in enumerate(items, 1):
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(f"{idx}. {text}")
            run.bold = False
            run.font.size = Pt(11)
        # Párrafo siguiente
        texto4 = (
            "Por los motivos expuestos les reitero mi propuesta de prestarles servicios independientes; es decir "
            "sin subordinación, dirección u horarios específicos y que de conformidad con lo dispuesto en el Artículo 94 fracción V de la Ley  del Impuesto Sobre la Renta, solicito que los ingresos que percibo de ustedes por la prestación de servicios profesionales independientes se asimilen a salarios y me retengan el Impuesto sobre la Renta de conformidad con lo dispuesto por el artículo 96 de la Ley mencionada. Para tal fin adjunto copia de la constancia de mi registro en el RFC del Servicio de Administración Tributaria."
        )
        p = document.add_paragraph()
        run = p.add_run(texto4)
        run.bold = False
        run.font.size = Pt(11)
        
        # Párrafo final
        texto5 = (
            "Asimismo, manifiesto que esta solicitud no dará lugar a ningún cambio en la relación contractual, ni "
            "tampoco generará ninguna relación laboral, por lo que eximo a esa empresa de cualquier responsabilidad laboral o legal que pudiera surgir."
        )
        p = document.add_paragraph()
        run = p.add_run(texto5)
        run.bold = False
        run.font.size = Pt(11)
        
        # Despedida
        p = document.add_paragraph()
        run = p.add_run("No habiendo otro asunto que tratar y esperando verme favorecido en esta petición, agradezco de antemano sus atenciones.")
        run.bold = False
        run.font.size = Pt(11)
        # document.add_paragraph()
        p = document.add_paragraph()
        run = p.add_run("Atentamente")
        run.bold = True
        run.font.size = Pt(11)  
        # Firma (nombre y RFC)
        p = document.add_paragraph()
        run = p.add_run(partes.nombreParte)
        run.bold = True
        run.font.size = Pt(11)
        
        p = document.add_paragraph()
        run = p.add_run("RFC: " + partes.rfc)
        run.bold = True
        run.font.size = Pt(11)
    
    
    

    document_data = io.BytesIO()
    document.save(document_data)
    document_data.seek(0)
    response = HttpResponse(
        document_data.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = 'attachment; filename="Contrato.docx"'
    response["Content-Encoding"] = "UTF-8"
    return response


@login_required(login_url="/login/")
@permission_required("cto.change_contratos", login_url="/login/")
def contratosAvanza(request, id):
    contratos = Contratos.objects.filter(pk=id).first()

    if request.method == "POST":

        if not contratos.fcap:
            contratos.fcap = datetime.today()
            contratos.current_user = contratos.rstep1
            contratos.status = contratos.astep1
            contratos.devuelto_por = contratos.rcap
            contratos.save()
            return HttpResponse("OK")
        # return HttpResponse("FAIL")
        if not contratos.fstep1:
            contratos.fstep1 = datetime.today()
            contratos.current_user = contratos.rstep2
            contratos.status = contratos.astep2
            contratos.save()
            return HttpResponse("OK")
        if not contratos.fstep2:
            contratos.fstep2 = datetime.today()
            contratos.current_user = contratos.rstep3
            contratos.status = contratos.astep3
            contratos.save()
            return HttpResponse("OK")
        if not contratos.fstep3:
            contratos.fstep3 = datetime.today()
            contratos.current_user = contratos.rstep4
            contratos.status = contratos.astep4
            contratos.save()
            return HttpResponse("OK")
        if not contratos.fstep4:
            contratos.fstep4 = datetime.today()
            contratos.current_user = contratos.rstep5
            contratos.status = contratos.astep5
            contratos.save()
            return HttpResponse("OK")
        if not contratos.fstep5:
            contratos.fstep5 = datetime.today()
            contratos.current_user = contratos.rstep6
            contratos.status = contratos.astep6
            contratos.save()
            return HttpResponse("OK")

        if not contratos.fstep6:
            contratos.fstep6 = datetime.today()
            #contratos.current_user = contratos.rstep6
            contratos.status = "FIN"
            contratos.save()
            return HttpResponse("OK")
        return HttpResponse("FAIL")

    return HttpResponse("FAIL")


@login_required(login_url="/login/")
@permission_required("cto.change_contratos", login_url="/login/")
def contratosDevuelve(request, id):
    contratos = Contratos.objects.filter(pk=id).first()

    if request.method == "POST":
        if not contratos.fstep6:
            contratos.fcap = contratos.fstep6
            contratos.fstep5 = contratos.fstep6
            contratos.fstep4 = contratos.fstep6
            contratos.fstep3 = contratos.fstep6
            contratos.fstep2 = contratos.fstep6
            contratos.fstep1 = contratos.fstep6
            contratos.devuelto_por = contratos.current_user
            contratos.current_user = contratos.rcap
            contratos.status = "CAP"

            contratos.save()
            return HttpResponse("OK")
        return HttpResponse("FAIL")
    return HttpResponse("FAIL")


def numero_to_letras(numero):

    indicador = [
        ("", ""), ("MIL", "MIL"), ("MILLON", "MILLONES"), ("MIL", "MIL"), ("BILLON", "BILLONES")]

    entero = int(numero)

    decimal = int(round((numero - entero)*100))

    contador = 0

    numero_letras = ""

    while entero > 0:

        a = entero % 1000

        if contador == 0:

            en_letras = convierte_cifra(a, 1).strip()

        else:

            en_letras = convierte_cifra(a, 0).strip()

        if a == 0:

            numero_letras = en_letras+" "+numero_letras

        elif a == 1:

            if contador in (1, 3):

                numero_letras = indicador[contador][0]+" "+numero_letras

            else:

                numero_letras = en_letras+" " + \
                    indicador[contador][0]+" "+numero_letras

        else:

            numero_letras = en_letras+" " + \
                indicador[contador][1]+" "+numero_letras

        numero_letras = numero_letras.strip()

        contador = contador + 1

        entero = int(entero / 1000)

    numero_letras = numero_letras+" PESOS " + str(decimal) + "/100 M.N."
    return (numero_letras)


def convierte_cifra(numero, sw):

    lista_centana = ["", ("CIEN", "CIENTO"), "DOSCIENTOS", "TRESCIENTOS", "CUATROCIENTOS",
                     "QUINIENTOS", "SEISCIENTOS", "SETECIENTOS", "OCHOCIENTOS", "NOVECIENTOS"]

    lista_decena = ["", ("DIEZ", "ONCE", "DOCE", "TRECE", "CATORCE", "QUINCE", "DIECISEIS", "DIECISIETE", "DIECIOCHO", "DIECINUEVE"),
                    ("VEINTE", "VEINTI"), ("TREINTA",
                                           "TREINTA Y"), ("CUARENTA", "CUARENTA Y"),
                    ("CINCUENTA", "CINCUENTA Y"), ("SESENTA", "SESENTA Y"),
                    ("SETENTA", "SETENTA Y"), ("OCHENTA", "OCHENTA Y"),
                    ("NOVENTA", "NOVENTA Y")
                    ]

    lista_unidad = ["", ("UN", "UNO"), "DOS", "TRES",
                    "CUATRO", "CINCO", "SEIS", "SIETE", "OCHO", "NUEVE"]

    centena = int(numero / 100)

    decena = int((numero - (centena * 100))/10)

    unidad = int(numero - (centena * 100 + decena * 10))

    texto_centena = ""

    texto_decena = ""

    texto_unidad = ""

    texto_centena = lista_centana[centena]

    if centena == 1:

        if (decena + unidad) != 0:

            texto_centena = texto_centena[1]

        else:

            texto_centena = texto_centena[0]

    texto_decena = lista_decena[decena]

    if decena == 1:

        texto_decena = texto_decena[unidad]

    elif decena > 1:

        if unidad != 0:

            texto_decena = texto_decena[1]

        else:

            texto_decena = texto_decena[0]

    if decena != 1:

        texto_unidad = lista_unidad[unidad]

        if unidad == 1:

            texto_unidad = texto_unidad[sw]

    return "%s %s %s" % (texto_centena, texto_decena, texto_unidad)


def _add_field(run, field):
    """ add a field to a run
    """
    fldChar1 = OxmlElement('w:fldChar')  # creates a new element
    fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = field

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "Seq"
    fldChar2.append(t)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)


def _add_number_range(run):
    """ add a number range field to a run
    """
    _add_field(run, r'Page')


@login_required(login_url="/login/")
@permission_required("cto.change_contratos", login_url="/login/")
def contratoGracont(request, id):
    contratos = Contratos.objects.filter(pk=id).first()

    if request.is_ajax and request.method == "POST":

        data = json.loads(request.body)
        # print(data)

        if not contratos.fcap:
            contratos.datecontrato = datetime.today()
            contratos.datecontrato_ini = data["enc_datecontrato_ini"]
            contratos.datecontrato_fin = data["enc_datecontrato_fin"]
            xiC = data["enc_importeContrato"]
            # print(xiC)
            # print(type(xiC))
            #contratos.importeContrato = data["enc_importeContrato"]

            contratos.npContrato = data["enc_npContrato"]

            contratos.empresa = data["enc_empresa"]
            contratos.tipo_sociedad = data["enc_tipo_sociedad"]
            contratos.objeto_social = data["enc_objeto_social"]
            contratos.registro_mercantil = data["enc_registro_mercantil"]
            contratos.fecha_constitucion = data["enc_fecha_constitucion"]
            contratos.domicilio_sociedad = data["enc_domicilio_sociedad"]
            contratos.rfc_sociedad = data["enc_rfc_sociedad"]
            contratos.testimonio = data["enc_testimonio"]
            contratos.clausula = data["enc_clausula"]
            
                
            
                     
            
            
            
            contratos.marcaVehiculo = data["enc_marcaVehiculo"]
            contratos.modeloVehiculo = data["enc_modeloVehiculo"]
            contratos.tipoVehiculo = data["enc_tipoVehiculo"]
            contratos.motorVehiculo = data["enc_motorVehiculo"]
            contratos.serieVehiculo = data["enc_serieVehiculo"]
            contratos.placasVehiculo = data["enc_placasVehiculo"]
            contratos.facturaVehiculo = data["enc_facturaVehiculo"]
            contratos.fechaFactura = data["enc_fechaFactura"]
            contratos.expedidaFactura = data["enc_expedidaFactura"]
            contratos.tarjetaVehiculo = data["enc_tarjetaVehiculo"]
            contratos.fechaTarjeta = data["enc_fechaTarjeta"]
            contratos.expedidaTarjeta = data["enc_expedidaTarjeta"]
            contratos.polizaVehiculo = data["enc_polizaVehiculo"]
            contratos.fechaPoliza = data["enc_fechaPoliza"]
            contratos.expedidaPoliza = data["enc_expedidaPoliza"]

            #contratos.imppContrato = data["enc_imppContrato"]
            contratos.totalhorasContrato = data["enc_totalhorasContrato"]
            contratos.testigoContrato1 = data["enc_testigoContrato1"]
            contratos.testigoContrato2 = data["enc_testigoContrato2"]

            contratos.save()
            return HttpResponse("OK")

        return HttpResponse("FAIL")

    return HttpResponse("FAIL")


@login_required(login_url="/login/")
@permission_required("cto.change_contratos", login_url="/login/")
def marcaContrato(request, id):
    tipocontratos = Tipocontrato.objects.filter(pk=id).first()

    if request.is_ajax and request.method == "POST":

        data = json.loads(request.body)
        # print(data)

        if Tipocontrato.tipoContrato:
            Tipocontrato.marcatipoContrato = True
            Tipocontrato.save()
            return HttpResponse("OK")

        return HttpResponse("FAIL")

    return HttpResponse("FAIL")


def calcular_edad_anos(fecha_nacimiento):
    fecha_actual = date.today()
    resultado = fecha_actual.year - fecha_nacimiento.year
    resultado -= ((fecha_actual.month, fecha_actual.day) <
                  (fecha_nacimiento.month, fecha_nacimiento.day))
    return resultado


class DoctosDetDelete(SinPrivilegios, generic.DeleteView):
    permission_required = "cto.delete_doctos"
    model = Doctos
    template_name = "cto/doctos_det_del.html"
    context_object_name = 'obj'

    def get_success_url(self):
        contrato_id = self.kwargs['contrato_id']
        return reverse_lazy('cto:contrato_edit', kwargs={'contrato_id': contrato_id})


class PuestosView(SinPrivilegios, generic.ListView):
    model = Puestos
    template_name = "cto/puestos_list.html"
    context_object_name = "obj"
    success_url = reverse_lazy("cto:puestos_list")
    permission_required = "cto.view_puestos"


class PuestosNew(VistaBaseCreate):
    model = Puestos
    template_name = "cto/puestos_form.html"
    form_class = PuestosForm
    success_url = reverse_lazy("cto:puestos_list")
    permission_required = "cto.add_puestos"


class PuestosEdit(VistaBaseEdit):
    model = Puestos
    template_name = "cto/puestos_form.html"
    form_class = PuestosForm
    success_url = reverse_lazy("cto:puestos_list")
    permission_required = "cto.change_puestos"


import requests

import json

def send_telegram_message(request, data):
    bot_token = '6316160840:AAFL0BJ828wpDz6RNlbpWE6Qv_GLEpkxx_w'
    chat_id = '-1002071180471'

    # Formatea el diccionario 'data' como una cadena de texto con saltos de línea y tabulaciones
    data_str = '\n'.join(f'{key}: {value}' for key, value in data.items())
    text = '*** Nuevo Contrato ***:\n' + data_str

    send_text = 'https://api.telegram.org/bot' + bot_token + '/sendMessage?chat_id=' + chat_id + '&parse_mode=Markdown&text=' + text

    response = requests.get(send_text)

    return HttpResponse('Mensaje enviado con éxito, respuesta: {}'.format(response.json()))
